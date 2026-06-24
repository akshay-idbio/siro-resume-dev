
"""
Low Cost Resume Requirement Matcher  FastAPI
============================================

Goal:
- Separate low-cost API script. Does NOT change your main app.
- Python parses resumes locally.
- Python prefilters top requirements.
- Claude Haiku gets only compact candidate + compact top requirements.
- 1 Claude call per resume.
- Output keeps main-app style sheets:
  1) Output
  2) Tracker
  3) Pivot
  4) Token_Cost

Run:
    pip install -r requirements_lowcost.txt
    uvicorn low_cost_resume_api:app --host 0.0.0.0 --port 8008 --reload

Endpoints:
    GET  /health
    POST /upload-requirement
    POST /bulk-analyze
    GET  /status/{job_id}
    GET  /download/{filename}
"""

import os
import re
import io
import json
import time
import uuid
import math
import shutil
import tempfile
import subprocess
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed

import pandas as pd
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from anthropic import Anthropic, APIStatusError

from pypdf import PdfReader
from docx import Document as DocxDocument

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from config import Settings, get_settings


# =========================================================
# LOW COST CONFIG
# =========================================================

APP_TITLE = "Low Cost Resume Matcher API"
APP_VERSION = "1.0.0"

BASE_DIR = Path(__file__).parent
WORK_DIR = BASE_DIR / "lowcost_workdir"
INPUT_DIR = WORK_DIR / "input"
RESUME_DIR = WORK_DIR / "resumes"
OUTPUT_DIR = WORK_DIR / "output"

for d in [WORK_DIR, INPUT_DIR, RESUME_DIR, OUTPUT_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# =========================================================
# MAIN CONFIG + LOW COST CONFIG
# =========================================================
# This low-cost app uses the SAME config.py and .env as your main app.
# API key comes from settings.anthropic_api_key.
# Low-cost model/settings come from the lowcost_* fields in config.py.

settings = get_settings()

LOW_COST_MODEL = settings.lowcost_claude_model
UPLOADED_REQUIREMENT_PATH = Path(settings.lowcost_requirement_excel_path)

MAX_PARALLEL = settings.lowcost_max_parallel

# This is the most important cost control.
TOP_REQUIREMENTS_FOR_CLAUDE = settings.lowcost_top_requirements
MAX_MATCHES_PER_RESUME = settings.lowcost_max_matches_per_resume
MAX_RESUME_CHARS_FOR_CLAUDE = settings.lowcost_max_resume_chars
MAX_JD_CHARS_FOR_CLAUDE = settings.lowcost_max_jd_chars
MAX_OUTPUT_TOKENS = settings.lowcost_max_output_tokens

# Update in config.py / .env if billing changes.
INPUT_RATE_PER_MILLION_USD = settings.lowcost_input_rate_per_million_usd
OUTPUT_RATE_PER_MILLION_USD = settings.lowcost_output_rate_per_million_usd
USD_TO_INR = settings.lowcost_usd_to_inr

SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx", ".doc"}

OUTPUT_COLUMNS = [
    "Request-ID",
    "MSP Owner",
    "Job Title",
    "Skills - Name",
    "Skills - Experience",
    "Additional Skills",
    "Job Description",
    "Work Location CDF",
    "Rate Card",
    "Annually",
    "CV File Name",
    "Candidate Name",
    "Candidate Phone",
    "Candidate Email",
    "Candidate Location",
    "Candidate Total Experience",
    "Candidate Skills",
    "Experience Mismatch",
    "Skill Mismatch",
    "ATS",
    "Remark",
]

TRACKER_COLUMNS = [
    "Date",
    "Request ID",
    "Status",
    "Skills",
    "Candidate",
    "Verdict",
    "Call Status",
    "Remarks",
]

PIVOT_REMARK_ORDER = [
    "Exp mismatch",
    "High CTC",
    "High Notice Period",
    "Location Mismatch",
    "Match",
    "Not Suitable",
    "NPU",
    "Skill mismatch",
]

VALID_REMARKS = set(PIVOT_REMARK_ORDER)
VALID_VERDICTS = {"Strong Fit", "Good Fit", "Possible Fit", "Not Suitable"}

VERDICT_COLORS = {
    "Strong Fit": "00B050",
    "Good Fit": "92D050",
    "Possible Fit": "FFC000",
    "Not Suitable": "FF6666",
}


# Expand this over time. Python extraction quality improves as this grows.
SKILL_VOCAB = [
    # cloud / devops
    "aws", "azure", "gcp", "kubernetes", "aks", "eks", "gke", "docker", "terraform",
    "ansible", "jenkins", "github actions", "gitlab", "ci/cd", "helm", "prometheus",
    "grafana", "datadog", "splunk", "linux", "rhel", "ubuntu", "bash", "shell",
    "devops", "sre", "iac", "cloudformation",

    # programming
    "python", "java", "javascript", "typescript", "c#", "c++", ".net", "golang",
    "go", "scala", "ruby", "php", "swift", "kotlin", "rust", "perl", "flask",
    "fastapi", "django", "spring", "spring boot", "microservices",

    # frontend
    "react", "react js", "reactjs", "angular", "vue", "nextjs", "next.js", "node",
    "node js", "nodejs", "express", "html", "css", "tailwind", "bootstrap",
    "spartacus",

    # databases / data
    "sql", "oracle", "postgres", "postgresql", "mysql", "mongodb", "cassandra",
    "redis", "snowflake", "databricks", "spark", "pyspark", "hadoop", "hive",
    "kafka", "airflow", "dbt", "tableau", "power bi", "qlik", "informatica",
    "iics", "teradata", "etl", "data warehouse", "data lake", "bigquery",

    # SAP / enterprise
    "sap", "s/4hana", "s4hana", "vistex", "hana", "abap", "sap mm", "sap sd",
    "sap fico", "sap crm", "sap bw", "datasphere", "sap hybris", "sap commerce",
    "ariba", "successfactors",

    # AI / ML
    "machine learning", "ml", "deep learning", "tensorflow", "pytorch", "llm",
    "nlp", "genai", "generative ai", "scikit", "scikit-learn", "pandas",
    "numpy", "opencv", "computer vision", "langchain", "llamaindex",

    # QA
    "selenium", "cypress", "jmeter", "qa", "automation testing", "manual testing",
    "appium", "playwright",

    # business apps
    "workday", "salesforce", "servicenow", "oracle fusion", "netsuite",

    # pharma / regulatory
    "gxp", "gmp", "clinical", "cdisc", "sdtm", "rbqm", "etmf", "ich",
    "21 cfr part 11", "veeva", "pharma", "pharmacovigilance", "regulatory",

    # project / process
    "jira", "confluence", "git", "agile", "scrum", "itil", "kanban",
    "nexthink", "citrix", "vmware", "sccm", "active directory",
]



# =========================================================
# APP
# =========================================================

app = FastAPI(title=APP_TITLE, version=APP_VERSION)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # restrict this in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

JOB_STATE: Dict[str, Dict[str, Any]] = {}


# =========================================================
# MODELS
# =========================================================

class UploadRequirementResponse(BaseModel):
    message: str
    filename: str
    total_requirements: int
    columns: List[str]


class BulkAnalyzeResponse(BaseModel):
    message: str
    job_id: str
    status_url: str


# =========================================================
# BASIC HELPERS
# =========================================================

def debug(msg: str) -> None:
    print(f"[LOWCOST DEBUG] {msg}", flush=True)


def now_date_str() -> str:
    return time.strftime("%Y-%m-%d")


def clean_cell(value: Any) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    value = str(value)
    value = value.replace("_x000D_", "\n").replace("\r", "\n")
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def normalize_text(value: Any) -> str:
    value = clean_cell(value).lower()
    replacements = {
        "reactjs": "react js",
        "nodejs": "node js",
        "next.js": "nextjs",
        "postgresql": "postgres sql",
        "snowflakes": "snowflake",
        "powerplatform": "power platform",
        "s/4 hana": "s4hana",
        "s/4hana": "s4hana",
    }
    for a, b in replacements.items():
        value = value.replace(a, b)
    value = re.sub(r"[^a-z0-9+#.\s/-]", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def safe_number(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        if isinstance(value, str):
            value = value.replace(",", "").strip()
        return float(value)
    except Exception:
        return default


def clamp_int(value: Any, min_value: int = 0, max_value: int = 100) -> int:
    try:
        v = int(round(float(value)))
    except Exception:
        v = 0
    return max(min_value, min(max_value, v))


def split_skills(value: Any) -> List[str]:
    value = clean_cell(value)
    parts = re.split(r"[,~|;/\n]+", value)
    return [p.strip() for p in parts if p.strip()]


def cut_text(value: Any, max_chars: int) -> str:
    value = clean_cell(value)
    if len(value) <= max_chars:
        return value
    return value[:max_chars].rsplit(" ", 1)[0].strip()


def clean_json_text(raw: str) -> dict:
    clean = raw.strip()
    clean = re.sub(r"^```(?:json)?\s*", "", clean)
    clean = re.sub(r"\s*```$", "", clean)
    start = clean.find("{")
    end = clean.rfind("}")
    if start != -1 and end != -1:
        clean = clean[start:end + 1]
    return json.loads(clean)


def verdict_from_ats(ats: Any) -> str:
    ats = safe_number(ats)
    if ats >= 85:
        return "Strong Fit"
    if ats >= 70:
        return "Good Fit"
    if ats >= 50:
        return "Possible Fit"
    return "Not Suitable"


def token_cost_usd(input_tokens: int, output_tokens: int) -> float:
    return (
        (input_tokens / 1_000_000.0) * INPUT_RATE_PER_MILLION_USD
        + (output_tokens / 1_000_000.0) * OUTPUT_RATE_PER_MILLION_USD
    )


def set_job(job_id: str, **kwargs: Any) -> None:
    st = JOB_STATE.setdefault(job_id, {})
    st.update(kwargs)
    st["updated_at"] = time.time()


def add_job_log(job_id: str, message: str) -> None:
    st = JOB_STATE.setdefault(job_id, {})
    logs = st.setdefault("log", [])
    logs.append({"ts": time.strftime("%H:%M:%S"), "message": message})
    st["log"] = logs[-100:]
    st["message"] = message
    st["updated_at"] = time.time()
    debug(f"{job_id}: {message}")


# =========================================================
# REQUIREMENT LOADING + COMPACTING
# =========================================================

def normalize_requirement_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    rename_map = {
        "Request ID": "Request-ID",
        "Request Id": "Request-ID",
        "Req ID": "Request-ID",
        "Name of Skills": "Skills - Name",
        "Skills": "Skills - Name",
        "Job Position": "Job Title",
        "Position": "Job Title",
        "Name w/o Link": "MSP Owner",
        "SPOC": "MSP Owner",
        "Monthly Company Pay Rate": "Rate Card",
        "Monthly Company Pay Rate ": "Rate Card",
        "Anually Company Pay Rate": "Yearly Rate",
        "Annually Company Pay Rate": "Yearly Rate",
        "Annual Company Pay Rate": "Yearly Rate",
        "Annually": "Yearly Rate",
        "Work Location City": "Work Location City",
        "Work Location CDF": "Work Location CDF",
    }
    df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})

    required_min = ["Request-ID"]
    missing = [c for c in required_min if c not in df.columns]
    if missing:
        raise HTTPException(status_code=400, detail=f"Requirement Excel missing: {missing}")

    defaults = {
        "MSP Owner": "",
        "Job Title": "",
        "Skills - Name": "",
        "Skills - Experience": "",
        "Additional Skills": "",
        "Job Description": "",
        "Status": "",
        "Work Location City": "",
        "Work Location CDF": "",
        "Rate Card": "",
        "Yearly Rate": "",
        "System Enhancements Required": "",
        "Candidate Annual CTC": "",
    }
    for c, default in defaults.items():
        if c not in df.columns:
            df[c] = default

    df["Request-ID"] = df["Request-ID"].astype(str).str.strip()
    df = df[df["Request-ID"] != ""]
    df = df.dropna(how="all")
    return df


def load_requirement_df() -> pd.DataFrame:
    if not UPLOADED_REQUIREMENT_PATH.exists():
        raise HTTPException(
            status_code=400,
            detail=(
                "Requirement Excel not found. Please upload Requirement Excel from the Main page first. "
                f"Expected path: {UPLOADED_REQUIREMENT_PATH}"
            ),
        )

    xls = pd.ExcelFile(UPLOADED_REQUIREMENT_PATH)
    matched = [s for s in xls.sheet_names if "requirement" in str(s).lower()]
    sheet_name = matched[0] if matched else xls.sheet_names[0]

    df = pd.read_excel(UPLOADED_REQUIREMENT_PATH, sheet_name=sheet_name)
    df = normalize_requirement_columns(df)

    if df.empty:
        raise HTTPException(status_code=400, detail="Requirement Excel has no valid Request-ID rows.")

    return df


def requirement_for_output(row: Dict[str, Any]) -> Dict[str, str]:
    return {
        "Request-ID": clean_cell(row.get("Request-ID")),
        "MSP Owner": clean_cell(row.get("MSP Owner")),
        "Job Title": clean_cell(row.get("Job Title")),
        "Skills - Name": clean_cell(row.get("Skills - Name")),
        "Skills - Experience": clean_cell(row.get("Skills - Experience")),
        "Additional Skills": clean_cell(row.get("Additional Skills")),
        "Job Description": clean_cell(row.get("Job Description")),
        "Work Location CDF": clean_cell(row.get("Work Location CDF")),
        "Rate Card": clean_cell(row.get("Rate Card")),
        "Annually": clean_cell(row.get("Yearly Rate") or row.get("Annually")),
        "Status": clean_cell(row.get("Status")),
    }


def extract_requirement_keywords(row: Dict[str, Any]) -> List[str]:
    text = " ".join([
        clean_cell(row.get("Job Title")),
        clean_cell(row.get("Skills - Name")),
        clean_cell(row.get("Additional Skills")),
        clean_cell(row.get("Job Description")),
    ]).lower()

    found = []
    for skill in SKILL_VOCAB:
        s = skill.lower()
        if any(ch in s for ch in [" ", ".", "/", "+", "#", "-"]):
            if s in text:
                found.append(skill)
        else:
            if re.search(rf"\b{re.escape(s)}\b", text):
                found.append(skill)

    # Add primary skill strings too.
    for s in split_skills(row.get("Skills - Name"))[:20]:
        if s and s.lower() not in {x.lower() for x in found}:
            found.append(s)

    return found[:40]


def compact_requirement(row: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "request_id": clean_cell(row.get("Request-ID")),
        "job_title": cut_text(row.get("Job Title"), 120),
        "status": clean_cell(row.get("Status")),
        "must_have_skills": split_skills(row.get("Skills - Name"))[:30],
        "skills_experience": cut_text(row.get("Skills - Experience"), 300),
        "additional_skills": split_skills(row.get("Additional Skills"))[:30],
        "keywords_from_jd": extract_requirement_keywords(row)[:35],
        "experience_required": extract_required_experience(row),
        "location": clean_cell(row.get("Work Location CDF") or row.get("Work Location City")),
        "rate_card": clean_cell(row.get("Rate Card")),
        "yearly_rate": clean_cell(row.get("Yearly Rate")),
        "jd_compact": cut_text(row.get("Job Description"), MAX_JD_CHARS_FOR_CLAUDE),
    }


def extract_required_experience(row: Dict[str, Any]) -> str:
    candidate_text = " ".join([
        clean_cell(row.get("Skills - Experience")),
        clean_cell(row.get("Job Description")),
    ])
    m = re.search(r"(\d{1,2})\s*[-to]+\s*(\d{1,2})\s*(?:years|yrs|year|yr)", candidate_text, flags=re.I)
    if m:
        return f"{m.group(1)}-{m.group(2)} years"
    m = re.search(r"(\d{1,2})\s*\+?\s*(?:years|yrs|year|yr)", candidate_text, flags=re.I)
    if m:
        return f"{m.group(1)}+ years"
    return clean_cell(row.get("Skills - Experience"))[:100]


# =========================================================
# RESUME PARSING WITH PAGE + CHAR POSITIONS
# =========================================================

def parse_pdf_text_blocks(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    text_parts: List[str] = []
    blocks: List[Dict[str, Any]] = []
    cursor = 0

    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = page_text.replace("\x00", " ")
        start = cursor
        text_parts.append(page_text)
        cursor += len(page_text)
        end = cursor
        blocks.append({"page": i, "start_char": start, "end_char": end, "text": page_text[:2000]})
        text_parts.append("\n")
        cursor += 1

    full_text = "\n".join([b["text"] for b in blocks])
    # Recalculate blocks using full page text not only preview.
    full_parts = []
    blocks = []
    cursor = 0
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = re.sub(r"\s+\n", "\n", page_text).strip()
        start = cursor
        full_parts.append(page_text)
        cursor += len(page_text)
        end = cursor
        blocks.append({
            "page": i,
            "start_char": start,
            "end_char": end,
            "preview": page_text[:500],
        })
        full_parts.append("\n")
        cursor += 1

    return "\n".join(full_parts).strip(), blocks


def parse_docx_text_blocks(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    doc = DocxDocument(str(path))
    chunks: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            chunks.append(p.text)

    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " | ".join(clean_cell(c.text) for c in row.cells if clean_cell(c.text))
            if row_text:
                chunks.append(row_text)

    full_text = "\n".join(chunks).strip()
    blocks = []
    cursor = 0
    for idx, para in enumerate(chunks, start=1):
        start = cursor
        cursor += len(para)
        blocks.append({
            "page": None,
            "block": idx,
            "start_char": start,
            "end_char": cursor,
            "preview": para[:500],
        })
        cursor += 1
    return full_text, blocks


def convert_doc_to_docx_or_pdf_bytes(path: Path) -> Optional[Path]:
    """
    Optional fallback for .doc using LibreOffice if available.
    Returns converted docx/pdf path or None.
    """
    try:
        with tempfile.TemporaryDirectory() as td:
            outdir = Path(td)
            cmd = [
                "libreoffice",
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "docx",
                "--outdir",
                str(outdir),
                str(path),
            ]
            res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=120)
            if res.returncode == 0:
                files = list(outdir.glob("*.docx"))
                if files:
                    final_path = path.with_suffix(".converted.docx")
                    shutil.copy(files[0], final_path)
                    return final_path
    except Exception:
        return None
    return None


def parse_resume_file(path: Path) -> Optional[Dict[str, Any]]:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            text, blocks = parse_pdf_text_blocks(path)
        elif ext == ".docx":
            text, blocks = parse_docx_text_blocks(path)
        elif ext == ".doc":
            converted = convert_doc_to_docx_or_pdf_bytes(path)
            if converted and converted.exists():
                text, blocks = parse_docx_text_blocks(converted)
                try:
                    converted.unlink()
                except Exception:
                    pass
            else:
                return None
        else:
            return None

        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()

        if len(text) < 80:
            return None

        return build_candidate_profile(path.name, text, blocks)

    except Exception as e:
        debug(f"Resume parse failed: {path.name}: {e}")
        return None


def extract_email(text: str) -> str:
    m = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else ""


def extract_phone(text: str) -> str:
    candidates = re.findall(r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,5}\)?[-.\s]?)?\d{3,5}[-.\s]?\d{4,6}", text)
    cleaned = []
    for c in candidates:
        digits = re.sub(r"\D", "", c)
        if 10 <= len(digits) <= 13:
            cleaned.append(c.strip())
    return cleaned[0] if cleaned else ""


def extract_name(text: str) -> str:
    labels = (
        "email", "phone", "mobile", "location", "address", "contact", "linkedin",
        "summary", "profile", "experience", "skills", "resume", "curriculum vitae",
        "objective", "education"
    )
    lines = [clean_cell(x) for x in text.splitlines() if clean_cell(x)]
    for line in lines[:15]:
        low = line.lower()
        if any(low.startswith(x) for x in labels):
            continue
        if ":" in line[:25]:
            continue
        if re.search(r"[\d@/|]", line):
            continue
        words = line.split()
        if 2 <= len(words) <= 5 and len(line) <= 60:
            return line
    return ""


def extract_years(text: str) -> Optional[float]:
    tl = text.lower()
    patterns = [
        r"total\s+experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years|yrs|year|yr)",
        r"(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years|yrs|year|yr)\s+of\s+(?:total\s+)?experience",
        r"experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years|yrs|year|yr)",
        r"(\d{1,2}(?:\.\d+)?)\s*\+\s*(?:years|yrs)",
    ]
    for p in patterns:
        m = re.search(p, tl)
        if m:
            return safe_number(m.group(1), None)
    return None


def extract_location(text: str) -> Dict[str, Any]:
    """
    Conservative location extraction. Avoids using random work locations too aggressively.
    """
    lines = [clean_cell(x) for x in text.splitlines() if clean_cell(x)]
    cities = [
        "bangalore", "bengaluru", "hyderabad", "chennai", "mumbai", "pune",
        "noida", "gurgaon", "gurugram", "delhi", "new delhi", "kolkata",
        "kochi", "ahmedabad", "vadodara", "surat", "indore", "jaipur",
        "nagpur", "nashik", "thane", "lucknow", "bhopal", "coimbatore",
        "trivandrum", "thiruvananthapuram", "mysore", "mangalore"
    ]

    # Best evidence: first 20 lines / contact section.
    header = "\n".join(lines[:25]).lower()
    for c in cities:
        if re.search(rf"\b{re.escape(c)}\b", header):
            city = "Bengaluru" if c == "bangalore" else c.title()
            return {
                "candidate_location": city,
                "candidate_city": city,
                "candidate_state": "",
                "candidate_country": "India",
                "candidate_location_source": "header/contact",
                "candidate_location_confidence": "medium",
            }

    # Explicit labels.
    explicit_patterns = [
        r"(?:current\s+location|location|address|residence)\s*[:\-]\s*([A-Za-z ,.-]{3,80})",
        r"(?:preferred\s+location)\s*[:\-]\s*([A-Za-z ,.-]{3,80})",
    ]
    for p in explicit_patterns:
        m = re.search(p, text, flags=re.I)
        if m:
            val = m.group(1).strip().split("\n")[0][:80]
            city = ""
            for c in cities:
                if re.search(rf"\b{re.escape(c)}\b", val.lower()):
                    city = "Bengaluru" if c == "bangalore" else c.title()
                    break
            return {
                "candidate_location": val,
                "candidate_city": city,
                "candidate_state": "",
                "candidate_country": "India" if city else "",
                "candidate_location_source": "current location",
                "candidate_location_confidence": "high" if city else "low",
            }

    return {
        "candidate_location": "",
        "candidate_city": "",
        "candidate_state": "",
        "candidate_country": "",
        "candidate_location_source": "not mentioned",
        "candidate_location_confidence": "not_found",
    }


def extract_skills(text: str) -> List[str]:
    tl = normalize_text(text)
    found = set()
    for skill in SKILL_VOCAB:
        s = normalize_text(skill)
        if not s:
            continue
        if any(ch in s for ch in [" ", ".", "/", "+", "#", "-"]):
            if s in tl:
                found.add(skill)
        else:
            if re.search(rf"\b{re.escape(s)}\b", tl):
                found.add(skill)
    return sorted(found, key=lambda x: x.lower())


def extract_recent_job_title(text: str) -> str:
    patterns = [
        r"(?:current\s+designation|current\s+role|designation|role|job\s+title)\s*[:\-]\s*([^\n]{3,80})",
        r"(?:working\s+as|worked\s+as)\s+([A-Za-z0-9 .,+#/-]{3,80})",
    ]
    for p in patterns:
        m = re.search(p, text, flags=re.I)
        if m:
            return clean_cell(m.group(1))[:80]
    return ""


def extract_ctc_notice(text: str) -> Dict[str, str]:
    current_ctc = ""
    expected_ctc = ""
    notice = ""

    m = re.search(r"(?:current\s+ctc|ctc)\s*[:\-]?\s*([^\n]{1,50})", text, flags=re.I)
    if m:
        current_ctc = clean_cell(m.group(1))[:50]
    m = re.search(r"(?:expected\s+ctc|ectc)\s*[:\-]?\s*([^\n]{1,50})", text, flags=re.I)
    if m:
        expected_ctc = clean_cell(m.group(1))[:50]
    m = re.search(r"(?:notice\s+period|notice)\s*[:\-]?\s*([^\n]{1,50})", text, flags=re.I)
    if m:
        notice = clean_cell(m.group(1))[:50]

    return {
        "candidate_current_ctc": current_ctc,
        "candidate_expected_ctc": expected_ctc,
        "candidate_notice_period": notice,
    }


def important_resume_text(text: str, max_chars: int) -> str:
    """
    Cost saver: prioritize contact/header, skills, summary, experience/project sections.
    """
    lines = [clean_cell(x) for x in text.splitlines() if clean_cell(x)]
    header = "\n".join(lines[:35])

    section_keywords = [
        "summary", "profile", "skills", "technical skills", "experience",
        "work experience", "professional experience", "projects", "education",
        "certification", "certifications"
    ]

    chosen = [header]
    lower_lines = [l.lower() for l in lines]

    for idx, low in enumerate(lower_lines):
        if any(k in low for k in section_keywords):
            chunk = "\n".join(lines[idx: idx + 45])
            chosen.append(chunk)

    compact = "\n\n".join(chosen)
    compact = re.sub(r"\n{3,}", "\n\n", compact)
    compact = re.sub(r"[ \t]+", " ", compact).strip()

    # If section detection failed, fallback to first max chars.
    if len(compact) < 1000:
        compact = text[:max_chars]

    return cut_text(compact, max_chars)


def build_candidate_profile(filename: str, text: str, blocks: List[Dict[str, Any]]) -> Dict[str, Any]:
    loc = extract_location(text)
    ctc_notice = extract_ctc_notice(text)
    skills = extract_skills(text)
    exp = extract_years(text)
    profile_summary = important_resume_text(text, MAX_RESUME_CHARS_FOR_CLAUDE)

    return {
        "filename": filename,
        "candidate_name": extract_name(text),
        "candidate_phone": extract_phone(text),
        "candidate_email": extract_email(text),
        **loc,
        "candidate_total_experience_years": exp,
        **ctc_notice,
        "candidate_skills": skills,
        "recent_job_title": extract_recent_job_title(text),
        "profile_summary": profile_summary,
        "_full_text_chars": len(text),
        "_sent_text_chars": len(profile_summary),
        "_text_blocks": blocks[:20],
    }


# =========================================================
# PREFILTER REQUIREMENTS PER RESUME
# =========================================================

def build_requirement_search_text(row: Dict[str, Any]) -> str:
    return " ".join([
        clean_cell(row.get("Job Title")),
        clean_cell(row.get("Skills - Name")),
        clean_cell(row.get("Skills - Experience")),
        clean_cell(row.get("Additional Skills")),
        clean_cell(row.get("Job Description")),
        clean_cell(row.get("Work Location CDF")),
        clean_cell(row.get("Work Location City")),
    ])


def build_candidate_search_text(candidate: Dict[str, Any]) -> str:
    return " ".join([
        " ".join(candidate.get("candidate_skills") or []),
        clean_cell(candidate.get("recent_job_title")),
        clean_cell(candidate.get("profile_summary")),
        clean_cell(candidate.get("candidate_location")),
    ])


def experience_score(candidate_exp: Optional[float], req_text: str) -> float:
    if candidate_exp is None:
        return 0.2

    req_text = clean_cell(req_text)
    m = re.search(r"(\d{1,2})\s*[-to]+\s*(\d{1,2})\s*(?:years|yrs|year|yr)", req_text, flags=re.I)
    if m:
        lo, hi = safe_number(m.group(1)), safe_number(m.group(2))
        if lo <= candidate_exp <= hi:
            return 1.0
        if candidate_exp >= lo - 1 and candidate_exp <= hi + 2:
            return 0.7
        return 0.2

    m = re.search(r"(\d{1,2})\s*\+?\s*(?:years|yrs|year|yr)", req_text, flags=re.I)
    if m:
        lo = safe_number(m.group(1))
        if candidate_exp >= lo:
            return 1.0
        if candidate_exp >= lo - 1:
            return 0.6
        return 0.2

    return 0.3


def location_score(candidate: Dict[str, Any], row: Dict[str, Any]) -> float:
    cand_city = normalize_text(candidate.get("candidate_city") or candidate.get("candidate_location"))
    req_loc = normalize_text(row.get("Work Location CDF") or row.get("Work Location City"))
    conf = candidate.get("candidate_location_confidence")

    if not cand_city or conf in {"low", "not_found", None, ""}:
        return 0.3
    if cand_city and cand_city in req_loc:
        return 1.0
    if "remote" in req_loc or "pan india" in req_loc or "anywhere" in req_loc:
        return 0.8
    return 0.0


def prefilter_requirements(requirements_df: pd.DataFrame, candidate: Dict[str, Any], top_n: int) -> List[Dict[str, Any]]:
    candidate_text = build_candidate_search_text(candidate)
    candidate_norm = normalize_text(candidate_text)
    candidate_tokens = set(candidate_norm.split())
    candidate_skills_norm = {normalize_text(s) for s in candidate.get("candidate_skills", [])}

    scored: List[Tuple[float, Dict[str, Any]]] = []

    for _, row in requirements_df.iterrows():
        row_dict = row.fillna("").to_dict()
        req_text = build_requirement_search_text(row_dict)
        req_norm = normalize_text(req_text)
        req_tokens = set(req_norm.split())

        overlap_tokens = candidate_tokens & req_tokens
        token_score = min(len(overlap_tokens), 50) * 1.0

        primary_score = 0.0
        for skill in split_skills(row_dict.get("Skills - Name"))[:40]:
            sn = normalize_text(skill)
            if sn and (sn in candidate_norm or sn in candidate_skills_norm):
                primary_score += 12

        additional_score = 0.0
        for skill in split_skills(row_dict.get("Additional Skills"))[:40]:
            sn = normalize_text(skill)
            if sn and (sn in candidate_norm or sn in candidate_skills_norm):
                additional_score += 4

        title_score = 0.0
        title = normalize_text(row_dict.get("Job Title"))
        recent = normalize_text(candidate.get("recent_job_title"))
        if title and recent:
            title_tokens = set(title.split())
            recent_tokens = set(recent.split())
            title_score = len(title_tokens & recent_tokens) * 5

        exp_score = experience_score(candidate.get("candidate_total_experience_years"), req_text) * 10
        loc_score = location_score(candidate, row_dict) * 5

        status_score = 3 if normalize_text(row_dict.get("Status")) == "open" else 0

        total = token_score + primary_score + additional_score + title_score + exp_score + loc_score + status_score
        scored.append((total, row_dict))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [r for _, r in scored[:top_n]]


# =========================================================
# CLAUDE MATCHING  ONE CALL PER RESUME
# =========================================================

def get_anthropic_client() -> Anthropic:
    """
    Use the SAME Anthropic API key as the main app config.py / .env.
    No separate terminal ANTHROPIC_API_KEY is needed.
    """
    if not settings.anthropic_api_key:
        raise RuntimeError("Anthropic API key missing in main config.py / .env.")

    return Anthropic(api_key=settings.anthropic_api_key)


def prompt_candidate_for_claude(candidate: Dict[str, Any]) -> Dict[str, Any]:
    """
    Do not send full extracted blocks unless debugging.
    """
    return {
        "filename": candidate.get("filename"),
        "candidate_name": candidate.get("candidate_name") or "",
        "candidate_phone": candidate.get("candidate_phone") or "",
        "candidate_email": candidate.get("candidate_email") or "",
        "candidate_location": candidate.get("candidate_location") or "",
        "candidate_city": candidate.get("candidate_city") or "",
        "candidate_state": candidate.get("candidate_state") or "",
        "candidate_country": candidate.get("candidate_country") or "",
        "candidate_location_source": candidate.get("candidate_location_source") or "not mentioned",
        "candidate_location_confidence": candidate.get("candidate_location_confidence") or "not_found",
        "candidate_total_experience_years": candidate.get("candidate_total_experience_years"),
        "candidate_current_ctc": candidate.get("candidate_current_ctc") or "",
        "candidate_expected_ctc": candidate.get("candidate_expected_ctc") or "",
        "candidate_notice_period": candidate.get("candidate_notice_period") or "",
        "candidate_skills": (candidate.get("candidate_skills") or [])[:80],
        "recent_job_title": candidate.get("recent_job_title") or "",
        "profile_summary_compact": cut_text(candidate.get("profile_summary"), MAX_RESUME_CHARS_FOR_CLAUDE),
    }


def build_lowcost_prompt(candidate: Dict[str, Any], shortlisted: List[Dict[str, Any]]) -> str:
    compact_reqs = [compact_requirement(r) for r in shortlisted]
    candidate_compact = prompt_candidate_for_claude(candidate)

    return f"""
You are an expert technical recruiter.

You are given:
1. One candidate profile extracted locally from a resume.
2. A small shortlisted list of compact job requirements.

Your task:
- Evaluate the candidate against the shortlisted requirements only.
- Return maximum {MAX_MATCHES_PER_RESUME} relevant matches.
- A candidate can be mapped to multiple requirements when relevant.
- If no requirement is relevant, return exactly one closest requirement and mark it as Not Suitable.
- Do not invent Request-ID.
- Keep call_status as empty string "".

Important business rules:
- Relevant means meaningful skill, role, experience, or domain overlap.
- Do not return weak unrelated requirements just to increase rows.
- Use High CTC only if candidate CTC is available and exceeds budget.
- Use High Notice Period only if notice period is available and clearly high.
- Use Location Mismatch only when candidate location is explicitly available and conflicts with requirement location.
- If candidate location is missing or confidence is low/not_found, set location_mismatch = "Not Evaluated" and do not use final_remark = "Location Mismatch".
- Do not infer candidate location from employer, project, education, or client location.
- NPU means Not Picked Up; it cannot be detected from resume. Never set NPU here.

ATS guidance:
- 85-100 = Strong Fit
- 70-84 = Good Fit
- 50-69 = Possible Fit
- 0-49 = Not Suitable
- Score realistically based on core skills, secondary skills, experience, role/domain, location where available, and CTC/notice where available.

Allowed verdict:
- Strong Fit
- Good Fit
- Possible Fit
- Not Suitable

Allowed final_remark:
- Exp mismatch
- High CTC
- High Notice Period
- Location Mismatch
- Match
- Not Suitable
- Skill mismatch

Candidate Profile:
{json.dumps(candidate_compact, ensure_ascii=False)}

Shortlisted Requirements:
{json.dumps(compact_reqs, ensure_ascii=False)}

Return only valid JSON. No markdown.
Required JSON:
{{
  "candidate_corrections": {{
    "candidate_name": "",
    "candidate_phone": "",
    "candidate_email": "",
    "candidate_location": "",
    "candidate_total_experience_years": null,
    "candidate_skills": []
  }},
  "matches": [
    {{
      "request_id": "",
      "ats_score": 0,
      "verdict": "Not Suitable",
      "call_status": "",
      "final_remark": "Not Suitable",
      "experience_mismatch": "No",
      "skill_mismatch": "No",
      "location_mismatch": "Not Evaluated",
      "matching_skills": [],
      "missing_skills": [],
      "reason": ""
    }}
  ]
}}
"""


def call_claude_one_resume(candidate: Dict[str, Any], shortlisted: List[Dict[str, Any]]) -> Dict[str, Any]:
    client = get_anthropic_client()
    prompt = build_lowcost_prompt(candidate, shortlisted)

    try:
        resp = client.messages.create(
            model=LOW_COST_MODEL,
            max_tokens=MAX_OUTPUT_TOKENS,
            messages=[{"role": "user", "content": prompt}],
        )

        usage = getattr(resp, "usage", None)
        input_tokens = getattr(usage, "input_tokens", 0) if usage else 0
        output_tokens = getattr(usage, "output_tokens", 0) if usage else 0

        raw = resp.content[0].text.strip()
        result = clean_json_text(raw)

        result["_token_usage"] = {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": input_tokens + output_tokens,
            "cost_usd": token_cost_usd(input_tokens, output_tokens),
            "cost_inr": token_cost_usd(input_tokens, output_tokens) * USD_TO_INR,
        }
        return result

    except APIStatusError as e:
        raise RuntimeError(f"Claude API error {e.status_code}: {str(e)[:200]}")
    except Exception as e:
        raise RuntimeError(f"Claude matching failed: {str(e)[:300]}")


def validate_match_result(
    result: Dict[str, Any],
    candidate: Dict[str, Any],
    shortlisted: List[Dict[str, Any]],
) -> Dict[str, Any]:
    allowed_ids = {clean_cell(r.get("Request-ID")) for r in shortlisted}
    matches = result.get("matches")
    if not isinstance(matches, list):
        matches = []

    cleaned = []
    for m in matches:
        if not isinstance(m, dict):
            continue

        rid = clean_cell(m.get("request_id") or m.get("best_request_id"))
        if rid not in allowed_ids:
            continue

        ats = clamp_int(m.get("ats_score"), 0, 100)
        verdict = verdict_from_ats(ats)

        final_remark = clean_cell(m.get("final_remark"))
        if final_remark not in VALID_REMARKS:
            final_remark = "Not Suitable" if ats < 50 else "Match"

        # Safety: no location mismatch when location not reliable.
        loc_conf = candidate.get("candidate_location_confidence")
        cand_loc = clean_cell(candidate.get("candidate_location") or candidate.get("candidate_city"))
        if not cand_loc or loc_conf in {"low", "not_found", "", None}:
            if final_remark == "Location Mismatch":
                final_remark = "Not Suitable" if ats < 50 else "Match"
            location_mismatch = "Not Evaluated"
        else:
            location_mismatch = clean_cell(m.get("location_mismatch")) or "Not Evaluated"

        cleaned.append({
            "request_id": rid,
            "ats_score": ats,
            "verdict": verdict,
            "call_status": "",
            "final_remark": final_remark,
            "experience_mismatch": clean_cell(m.get("experience_mismatch")) or "No",
            "skill_mismatch": clean_cell(m.get("skill_mismatch")) or "No",
            "location_mismatch": location_mismatch,
            "matching_skills": m.get("matching_skills") if isinstance(m.get("matching_skills"), list) else [],
            "missing_skills": m.get("missing_skills") if isinstance(m.get("missing_skills"), list) else [],
            "reason": clean_cell(m.get("reason")),
        })

    # If Claude returns nothing valid, fallback to top prefilter as Not Suitable.
    if not cleaned and shortlisted:
        cleaned = [{
            "request_id": clean_cell(shortlisted[0].get("Request-ID")),
            "ats_score": 0,
            "verdict": "Not Suitable",
            "call_status": "",
            "final_remark": "Not Suitable",
            "experience_mismatch": "No",
            "skill_mismatch": "Yes",
            "location_mismatch": "Not Evaluated",
            "matching_skills": [],
            "missing_skills": [],
            "reason": "Fallback: Claude did not return a valid match. Closest prefiltered requirement selected.",
        }]

    # Selection rule:
    # If relevant exists, return all relevant up to MAX_MATCHES_PER_RESUME.
    # Else return closest/best one.
    cleaned.sort(key=lambda x: safe_number(x.get("ats_score")), reverse=True)
    relevant = [
        x for x in cleaned
        if safe_number(x.get("ats_score")) >= 50
        and x.get("verdict") != "Not Suitable"
        and x.get("final_remark") != "Not Suitable"
    ]
    if relevant:
        cleaned = relevant[:MAX_MATCHES_PER_RESUME]
    else:
        cleaned = cleaned[:1]

    result["matches"] = cleaned
    return result


def apply_candidate_corrections(candidate: Dict[str, Any], result: Dict[str, Any]) -> Dict[str, Any]:
    """
    Claude can correct candidate fields, but Python values remain fallback.
    """
    corr = result.get("candidate_corrections")
    if not isinstance(corr, dict):
        return candidate

    updated = dict(candidate)
    for key in [
        "candidate_name",
        "candidate_phone",
        "candidate_email",
        "candidate_location",
        "candidate_total_experience_years",
        "candidate_skills",
    ]:
        val = corr.get(key)
        if key == "candidate_skills":
            if isinstance(val, list) and val:
                updated[key] = val[:100]
        else:
            if val not in [None, "", "Not found in CV", "N/A"]:
                updated[key] = val
    return updated


# =========================================================
# OUTPUT BUILDING
# =========================================================

def build_rows_for_candidate(
    candidate: Dict[str, Any],
    result: Dict[str, Any],
    req_lookup: Dict[str, Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    output_rows = []
    tracker_rows = []

    skills_joined = ", ".join(candidate.get("candidate_skills") or [])

    for m in result.get("matches", []):
        rid = clean_cell(m.get("request_id"))
        req = req_lookup.get(rid, {})
        out_req = requirement_for_output(req)

        row = {
            **{c: out_req.get(c, "") for c in OUTPUT_COLUMNS},
            "CV File Name": candidate.get("filename", ""),
            "Candidate Name": candidate.get("candidate_name", ""),
            "Candidate Phone": candidate.get("candidate_phone", ""),
            "Candidate Email": candidate.get("candidate_email", ""),
            "Candidate Location": candidate.get("candidate_location", ""),
            "Candidate Total Experience": candidate.get("candidate_total_experience_years", ""),
            "Candidate Skills": skills_joined,
            "Experience Mismatch": m.get("experience_mismatch", "No"),
            "Skill Mismatch": m.get("skill_mismatch", "No"),
            "ATS": clamp_int(m.get("ats_score"), 0, 100),
            "Remark": m.get("final_remark", "Not Suitable"),
        }

        # Ensure Requirement-ID and Annually are correct after dict merge.
        row["Request-ID"] = out_req.get("Request-ID", rid)
        row["Annually"] = out_req.get("Annually", "")

        output_rows.append(row)

        tracker_rows.append({
            "Date": now_date_str(),
            "Request ID": rid,
            "Status": out_req.get("Status", ""),
            "Skills": out_req.get("Skills - Name", ""),
            "Candidate": candidate.get("candidate_name", "") or candidate.get("filename", ""),
            "Verdict": m.get("verdict", verdict_from_ats(m.get("ats_score"))),
            "Call Status": "",
            "Remarks": m.get("final_remark", "Not Suitable"),
        })

    return output_rows, tracker_rows


def autofit_sheet(ws, max_width: int = 55) -> None:
    for col_idx in range(1, ws.max_column + 1):
        col_letter = get_column_letter(col_idx)
        max_len = 0
        for cell in ws[col_letter]:
            value = "" if cell.value is None else str(cell.value)
            max_len = max(max_len, min(len(value), max_width))
        ws.column_dimensions[col_letter].width = max(12, min(max_len + 2, max_width))


def style_header(ws, fill_color: str = "0D2B4B") -> None:
    thin = Side(style="thin", color="D9E2F3")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", start_color=fill_color)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def write_rows(ws, headers: List[str], rows: List[Dict[str, Any]]) -> None:
    ws.append(headers)
    for r in rows:
        ws.append([r.get(h, "") for h in headers])
    style_header(ws)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    autofit_sheet(ws)


def build_pivot_rows(tracker_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    # Simple pivot: count remarks by Request ID.
    pivot_map: Dict[str, Dict[str, Any]] = {}
    for tr in tracker_rows:
        rid = clean_cell(tr.get("Request ID"))
        if not rid:
            continue
        row = pivot_map.setdefault(rid, {"Request ID": rid, **{k: 0 for k in PIVOT_REMARK_ORDER}, "Total": 0})
        remark = clean_cell(tr.get("Remarks"))
        if remark not in PIVOT_REMARK_ORDER:
            remark = "Not Suitable"
        row[remark] += 1
        row["Total"] += 1

    return list(pivot_map.values())


def build_excel_output(
    output_rows: List[Dict[str, Any]],
    tracker_rows: List[Dict[str, Any]],
    token_rows: List[Dict[str, Any]],
    filename: str,
) -> Path:
    wb = Workbook()

    ws_out = wb.active
    ws_out.title = "Output"
    write_rows(ws_out, OUTPUT_COLUMNS, output_rows)

    # Color ATS/verdict-ish remark cells.
    ats_idx = OUTPUT_COLUMNS.index("ATS") + 1
    remark_idx = OUTPUT_COLUMNS.index("Remark") + 1
    for r in range(2, ws_out.max_row + 1):
        ats = safe_number(ws_out.cell(r, ats_idx).value)
        fill = None
        if ats >= 85:
            fill = "00B050"
        elif ats >= 70:
            fill = "92D050"
        elif ats >= 50:
            fill = "FFC000"
        else:
            fill = "FF6666"
        ws_out.cell(r, ats_idx).fill = PatternFill("solid", start_color=fill)
        ws_out.cell(r, remark_idx).alignment = Alignment(vertical="top", wrap_text=True)

    ws_tracker = wb.create_sheet("Tracker")
    write_rows(ws_tracker, TRACKER_COLUMNS, tracker_rows)

    verdict_idx = TRACKER_COLUMNS.index("Verdict") + 1
    for r in range(2, ws_tracker.max_row + 1):
        v = ws_tracker.cell(r, verdict_idx).value
        color = VERDICT_COLORS.get(v)
        if color:
            ws_tracker.cell(r, verdict_idx).fill = PatternFill("solid", start_color=color)
            ws_tracker.cell(r, verdict_idx).font = Font(bold=True, color="FFFFFF")

    ws_pivot = wb.create_sheet("Pivot")
    pivot_headers = ["Request ID"] + PIVOT_REMARK_ORDER + ["Total"]
    write_rows(ws_pivot, pivot_headers, build_pivot_rows(tracker_rows))

    ws_tokens = wb.create_sheet("Token_Cost")
    token_headers = [
        "CV File Name", "Input Tokens", "Output Tokens", "Total Tokens",
        "Cost USD", "Cost INR", "Resume Full Chars", "Resume Sent Chars",
        "Top Requirements Sent", "Status", "Error"
    ]
    write_rows(ws_tokens, token_headers, token_rows)

    output_path = OUTPUT_DIR / filename
    wb.save(output_path)
    return output_path


# =========================================================
# JOB RUNNER
# =========================================================

def save_upload_file(upload: UploadFile, folder: Path) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    safe_name = Path(upload.filename or f"file_{uuid.uuid4().hex}").name
    dest = folder / safe_name
    with open(dest, "wb") as f:
        shutil.copyfileobj(upload.file, f)
    return dest


def process_single_resume(
    resume_path: Path,
    requirements_df: pd.DataFrame,
    req_lookup: Dict[str, Dict[str, Any]],
) -> Dict[str, Any]:
    token_row = {
        "CV File Name": resume_path.name,
        "Input Tokens": 0,
        "Output Tokens": 0,
        "Total Tokens": 0,
        "Cost USD": 0,
        "Cost INR": 0,
        "Resume Full Chars": 0,
        "Resume Sent Chars": 0,
        "Top Requirements Sent": 0,
        "Status": "Error",
        "Error": "",
    }

    parsed = parse_resume_file(resume_path)
    if not parsed:
        token_row["Error"] = "Could not parse resume or resume text too short."
        return {"output_rows": [], "tracker_rows": [], "token_row": token_row}

    shortlisted = prefilter_requirements(requirements_df, parsed, TOP_REQUIREMENTS_FOR_CLAUDE)
    token_row["Resume Full Chars"] = parsed.get("_full_text_chars", 0)
    token_row["Resume Sent Chars"] = parsed.get("_sent_text_chars", 0)
    token_row["Top Requirements Sent"] = len(shortlisted)

    try:
        raw_result = call_claude_one_resume(parsed, shortlisted)
        result = validate_match_result(raw_result, parsed, shortlisted)
        parsed = apply_candidate_corrections(parsed, result)

        usage = result.get("_token_usage") or raw_result.get("_token_usage") or {}
        token_row["Input Tokens"] = usage.get("input_tokens", 0)
        token_row["Output Tokens"] = usage.get("output_tokens", 0)
        token_row["Total Tokens"] = usage.get("total_tokens", 0)
        token_row["Cost USD"] = round(usage.get("cost_usd", 0), 6)
        token_row["Cost INR"] = round(usage.get("cost_inr", 0), 4)
        token_row["Status"] = "Success"
        token_row["Error"] = ""

        output_rows, tracker_rows = build_rows_for_candidate(parsed, result, req_lookup)
        return {
            "output_rows": output_rows,
            "tracker_rows": tracker_rows,
            "token_row": token_row,
        }

    except Exception as e:
        token_row["Error"] = str(e)[:300]
        # Still produce one tracking row? Safer to log only in Token_Cost.
        return {"output_rows": [], "tracker_rows": [], "token_row": token_row}


def run_bulk_job(job_id: str, resume_paths: List[Path]) -> None:
    start = time.time()
    try:
        set_job(job_id, status="running", stage="loading_requirements", progress=0, total=len(resume_paths))
        add_job_log(job_id, "Loading requirement Excel...")
        requirements_df = load_requirement_df()
        req_lookup = {
            clean_cell(row.get("Request-ID")): row.fillna("").to_dict()
            for _, row in requirements_df.iterrows()
        }

        output_rows: List[Dict[str, Any]] = []
        tracker_rows: List[Dict[str, Any]] = []
        token_rows: List[Dict[str, Any]] = []

        add_job_log(job_id, f"Processing {len(resume_paths)} resumes with max parallel {MAX_PARALLEL}...")
        done = 0

        with ThreadPoolExecutor(max_workers=MAX_PARALLEL) as pool:
            futures = {
                pool.submit(process_single_resume, p, requirements_df, req_lookup): p
                for p in resume_paths
            }
            for fut in as_completed(futures):
                p = futures[fut]
                try:
                    res = fut.result()
                    output_rows.extend(res.get("output_rows", []))
                    tracker_rows.extend(res.get("tracker_rows", []))
                    token_rows.append(res.get("token_row", {"CV File Name": p.name, "Status": "Error"}))
                except Exception as e:
                    token_rows.append({
                        "CV File Name": p.name,
                        "Input Tokens": 0,
                        "Output Tokens": 0,
                        "Total Tokens": 0,
                        "Cost USD": 0,
                        "Cost INR": 0,
                        "Resume Full Chars": 0,
                        "Resume Sent Chars": 0,
                        "Top Requirements Sent": 0,
                        "Status": "Error",
                        "Error": str(e)[:300],
                    })

                done += 1
                set_job(job_id, stage="matching", progress=done, total=len(resume_paths))
                if done % 5 == 0 or done == len(resume_paths):
                    add_job_log(job_id, f"Processed {done}/{len(resume_paths)} resumes...")

        if not output_rows:
            add_job_log(job_id, "No output rows created. Check Token_Cost sheet/errors.")

        ts = time.strftime("%Y%m%d_%H%M%S")
        out_name = f"lowcost_resume_requirement_output_{ts}_{job_id[:8]}.xlsx"
        out_path = build_excel_output(output_rows, tracker_rows, token_rows, out_name)

        total_cost_inr = sum(safe_number(r.get("Cost INR")) for r in token_rows)
        total_cost_usd = sum(safe_number(r.get("Cost USD")) for r in token_rows)
        total_tokens = sum(int(safe_number(r.get("Total Tokens"))) for r in token_rows)

        summary = {
            "total_files": len(resume_paths),
            "output_rows": len(output_rows),
            "tracker_rows": len(tracker_rows),
            "success_resumes": sum(1 for r in token_rows if r.get("Status") == "Success"),
            "error_resumes": sum(1 for r in token_rows if r.get("Status") != "Success"),
            "total_tokens": total_tokens,
            "total_cost_usd": round(total_cost_usd, 6),
            "total_cost_inr": round(total_cost_inr, 4),
            "avg_cost_inr_per_resume": round(total_cost_inr / max(len(resume_paths), 1), 4),
            "duration_seconds": round(time.time() - start, 2),
            "output_filename": out_name,
            "download_url": f"/download/{out_name}",
            "config": {
                "model": LOW_COST_MODEL,
        "api_key_source": "settings.anthropic_api_key from config.py",
                "top_requirements_for_claude": TOP_REQUIREMENTS_FOR_CLAUDE,
                "max_matches_per_resume": MAX_MATCHES_PER_RESUME,
                "max_resume_chars": MAX_RESUME_CHARS_FOR_CLAUDE,
                "max_jd_chars": MAX_JD_CHARS_FOR_CLAUDE,
                "max_parallel": MAX_PARALLEL,
            }
        }

        set_job(
            job_id,
            status="done",
            stage="done",
            progress=len(resume_paths),
            total=len(resume_paths),
            summary=summary,
            output_filename=out_name,
            download_url=f"/download/{out_name}",
        )
        add_job_log(job_id, f"Done. Output: {out_name}")

    except Exception as e:
        set_job(job_id, status="error", stage="error", error=str(e), summary=None)
        add_job_log(job_id, f"ERROR: {e}")


# =========================================================
# API ENDPOINTS
# =========================================================

@app.get("/health")
def health():
    return {
        "status": "ok",
        "app": APP_TITLE,
        "version": APP_VERSION,
        "model": LOW_COST_MODEL,
        "config": {
            "top_requirements_for_claude": TOP_REQUIREMENTS_FOR_CLAUDE,
            "max_matches_per_resume": MAX_MATCHES_PER_RESUME,
            "max_resume_chars": MAX_RESUME_CHARS_FOR_CLAUDE,
            "max_jd_chars": MAX_JD_CHARS_FOR_CLAUDE,
            "max_parallel": MAX_PARALLEL,
        }
    }


@app.post("/upload-requirement", response_model=UploadRequirementResponse)
async def upload_requirement(file: UploadFile = File(...)):
    filename = file.filename or ""
    if not filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Please upload Excel requirement file.")

    with open(UPLOADED_REQUIREMENT_PATH, "wb") as f:
        content = await file.read()
        f.write(content)

    df = load_requirement_df()
    return UploadRequirementResponse(
        message="Requirement Excel uploaded successfully.",
        filename=filename,
        total_requirements=len(df),
        columns=list(df.columns),
    )


@app.post("/bulk-analyze", response_model=BulkAnalyzeResponse)
async def bulk_analyze(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    if not UPLOADED_REQUIREMENT_PATH.exists():
        raise HTTPException(
            status_code=400,
            detail=(
                "Requirement Excel not found. Please upload Requirement Excel from the Main page first. "
                f"Expected path: {UPLOADED_REQUIREMENT_PATH}"
            ),
        )

    if not files:
        raise HTTPException(status_code=400, detail="No resume files uploaded.")

    job_id = uuid.uuid4().hex
    job_resume_dir = RESUME_DIR / job_id
    job_resume_dir.mkdir(parents=True, exist_ok=True)

    resume_paths: List[Path] = []
    skipped: List[str] = []

    for upload in files:
        ext = Path(upload.filename or "").suffix.lower()
        if ext not in SUPPORTED_RESUME_EXTENSIONS:
            skipped.append(upload.filename or "unknown")
            continue
        dest = save_upload_file(upload, job_resume_dir)
        resume_paths.append(dest)

    if not resume_paths:
        raise HTTPException(status_code=400, detail={"message": "No supported resume files.", "skipped": skipped})

    JOB_STATE[job_id] = {
        "job_id": job_id,
        "status": "queued",
        "stage": "queued",
        "progress": 0,
        "total": len(resume_paths),
        "message": "Queued",
        "log": [],
        "skipped_files": skipped,
        "created_at": time.time(),
    }

    background_tasks.add_task(run_bulk_job, job_id, resume_paths)

    return BulkAnalyzeResponse(
        message=f"Low-cost analysis started. {len(resume_paths)} resumes accepted, {len(skipped)} skipped.",
        job_id=job_id,
        status_url=f"/status/{job_id}",
    )


@app.get("/status/{job_id}")
def get_status(job_id: str):
    st = JOB_STATE.get(job_id)
    if not st:
        raise HTTPException(status_code=404, detail="Job ID not found.")
    return st


@app.get("/download/{filename}")
def download(filename: str):
    safe = Path(filename).name
    path = OUTPUT_DIR / safe
    if not path.exists():
        raise HTTPException(status_code=404, detail="Output file not found.")
    return FileResponse(
        path=str(path),
        filename=safe,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

# =========================================================
# LOW COST SEPARATE API ALIASES
# =========================================================
# Frontend LowCostAnalyze.jsx calls these endpoints.
# These aliases keep the low-cost pipeline separate and clear:
# POST /lowcost/bulk-analyze
# GET  /lowcost/status/{job_id}
# POST /lowcost/reset


@app.post("/lowcost/bulk-analyze", response_model=BulkAnalyzeResponse)
async def lowcost_bulk_analyze(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
):
    return await bulk_analyze(background_tasks, files)


@app.get("/lowcost/status/{job_id}")
def lowcost_get_status(job_id: str):
    return get_status(job_id)


@app.post("/lowcost/reset")
def lowcost_reset():
    finished_jobs = []

    for job_id, state in list(JOB_STATE.items()):
        if state.get("status") in ["done", "error", "failed", "completed"]:
            finished_jobs.append(job_id)
            JOB_STATE.pop(job_id, None)

    return {
        "message": "Low-cost finished job states cleared.",
        "cleared": len(finished_jobs),
    }
    
    
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("low_cost_resume_api:app", host="0.0.0.0", port=8008, reload=True)
