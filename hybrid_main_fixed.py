import os
import re
import json
import time
import uuid
import asyncio
import shutil
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Any

import anthropic
import pandas as pd
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from config import Settings, get_settings


# =========================================================
# HYBRID APP
# Goal:
# - Keep current main.py untouched.
# - Use local parsing + cheap AI matching.
# - Keep SAME business rule as main app:
#   If candidate matches multiple requirements => map to all relevant ones.
#   If candidate does not match anything => show only one closest/best fit.
# - Keep Excel fashion same as main app:
#   Output Sheet, Tracker, Pivot
# =========================================================

settings = get_settings()

app = FastAPI(
    title="Resume Requirement Hybrid Matching API",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # testing only; restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =========================================================
# CONSTANTS
# =========================================================

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

INPUT_DIR = "input"
os.makedirs(INPUT_DIR, exist_ok=True)

UPLOADED_REQUIREMENT_PATH = os.path.join(INPUT_DIR, "uploaded_requirement.xlsx")

HYBRID_JOB_DIR = "hybrid_jobs"
os.makedirs(HYBRID_JOB_DIR, exist_ok=True)

HYBRID_STATUS_PATH = os.path.join(HYBRID_JOB_DIR, "status.json")
HYBRID_UPLOAD_DIR = os.path.join(HYBRID_JOB_DIR, "resumes")
os.makedirs(HYBRID_UPLOAD_DIR, exist_ok=True)

POLLING_OUTPUT_PREFIX = "hybrid_resume_requirement_output"

# Use cheaper model for hybrid matching.
# Override using env var if required.
HYBRID_CLAUDE_MODEL = os.getenv("HYBRID_CLAUDE_MODEL", "claude-haiku-4-5-20251001")
HYBRID_MAX_TOKENS = int(os.getenv("HYBRID_MAX_TOKENS", "1800"))

# IMPORTANT:
# Same philosophy as main app. We should not shortlist too low,
# otherwise relevant requirements may be missed before Claude gets chance.
HYBRID_TOP_N_REQUIREMENTS = int(os.getenv("HYBRID_TOP_N_REQUIREMENTS", "25"))

# Process sequentially by default for stable status/cost tracking.
# Increase cautiously later.
HYBRID_MAX_PARALLEL = int(os.getenv("HYBRID_MAX_PARALLEL", "1"))

# Approx cost settings for Haiku. Update if billing changes.
HYBRID_INPUT_RATE_PER_MILLION_USD = float(os.getenv("HYBRID_INPUT_RATE_PER_MILLION_USD", "1.0"))
HYBRID_OUTPUT_RATE_PER_MILLION_USD = float(os.getenv("HYBRID_OUTPUT_RATE_PER_MILLION_USD", "5.0"))
USD_TO_INR = float(os.getenv("USD_TO_INR", "95.4"))

OUTPUT_COLUMNS = [
    "Request-ID",
    "MSP Owner",
    "Job Title",
    "Skills - Name",
    "Skills - Experience",
    "Additional Skills",
    "Job Description",
    "Work Location CDF",
    "Rate Card",
    "Annually",
    "CV File Name",
    "Candidate Name",
    "Candidate Phone",
    "Candidate Email",
    "Candidate Location",
    "Candidate Total Experience",
    "Candidate Skills",
    "Experience Mismatch",
    "Skill Mismatch",
    "ATS",
    "Remark",
]

REQUIRED_REQUIREMENT_COLUMNS = [
    "Request-ID",
    "MSP Owner",
    "Job Title",
    "Skills - Name",
    "Skills - Experience",
    "Additional Skills",
    "Job Description",
    "Status",
    "Work Location CDF",
    "Rate Card",
    "Yearly Rate",
]

TRACKER_COLUMNS = [
    "Date",
    "Request ID",
    "Status",
    "Skills",
    "Candidate",
    "Verdict",
    "Call Status",
    "Remarks",
]

PIVOT_REMARK_ORDER = [
    "Exp mismatch",
    "High CTC",
    "High Notice Period",
    "Location Mismatch",
    "Match",
    "Not Suitable",
    "NPU",
    "Skill mismatch",
]

VALID_REMARKS = set(PIVOT_REMARK_ORDER)

VERDICT_COLORS = {
    "Strong Fit": "00B050",
    "Good Fit": "92D050",
    "Possible Fit": "FFC000",
    "Not Suitable": "FF6666",
}


ALLOWED_RESUME_EXTENSIONS = [".pdf", ".docx", ".doc"]


# =========================================================
# RESPONSE MODEL
# =========================================================

class HybridStartResponse(BaseModel):
    job_id: str
    status: str
    total: int
    message: str


# =========================================================
# DEBUG
# =========================================================

def debug(msg: str):
    print(f"[HYBRID DEBUG] {msg}", flush=True)


# =========================================================
# BASIC CLEANING HELPERS
# =========================================================

def clean_cell(value) -> str:
    if value is None:
        return ""

    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass

    value = str(value)
    value = value.replace("_x000D_", "\n")
    value = value.replace("\r", "\n")
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def normalize_text(value) -> str:
    value = clean_cell(value).lower()
    value = value.replace("reactjs", "react js")
    value = value.replace("nodejs", "node js")
    value = value.replace("postgresql", "postgre sql")
    value = value.replace("snowflakes", "snowflake")
    value = value.replace("powerplatform", "power platform")
    value = re.sub(r"[^a-z0-9+#.\s]", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def split_skills(value) -> List[str]:
    value = clean_cell(value)
    parts = re.split(r"[,~|;/\n]+", value)
    return [p.strip() for p in parts if p.strip()]


def safe_number(value, default=0):
    try:
        if value is None or value == "":
            return default
        return float(value)
    except Exception:
        return default


def clamp_int(value, min_value=0, max_value=100) -> int:
    try:
        value = int(round(float(value)))
    except Exception:
        value = 0
    return max(min_value, min(max_value, value))


def clean_json_text(raw: str) -> dict:
    clean = raw.strip()

    if clean.startswith("```json"):
        clean = clean.removeprefix("```json").strip()

    if clean.startswith("```"):
        clean = clean.removeprefix("```").strip()

    if clean.endswith("```"):
        clean = clean.removesuffix("```").strip()

    start = clean.find("{")
    end = clean.rfind("}")

    if start != -1 and end != -1:
        clean = clean[start:end + 1]

    return json.loads(clean)


def get_file_extension(filename: str) -> str:
    filename = filename or ""
    return os.path.splitext(filename.lower())[1]


# =========================================================
# CORE BUSINESS RULE - SAME AS MAIN APP
# =========================================================

def select_relevant_matches_or_best_one(matches: list) -> list:
    """
    Business rule for Option 2 / multi-requirement mapping:

    - If candidate fits multiple jobs, return all relevant matches.
    - If candidate fits nothing, return only the closest/best one.
    - Do not show all Not Suitable rows.

    Relevance:
    - ATS >= 50
    - Verdict is not "Not Suitable"
    - final_remark is not "Not Suitable"

    This must be enforced in Python, not only in prompt,
    so output behavior remains stable.
    """

    if not matches:
        return []

    cleaned_matches = []

    for match in matches:
        if not isinstance(match, dict):
            continue

        ats = safe_number(match.get("ats_score"), 0)
        verdict = clean_cell(match.get("verdict")) or "Not Suitable"
        final_remark = clean_cell(match.get("final_remark")) or "Not Suitable"

        match["ats_score"] = ats
        match["verdict"] = verdict
        match["final_remark"] = final_remark
        match["call_status"] = ""

        cleaned_matches.append(match)

    if not cleaned_matches:
        return []

    cleaned_matches = sorted(
        cleaned_matches,
        key=lambda x: safe_number(x.get("ats_score"), 0),
        reverse=True,
    )

    relevant_matches = []

    for match in cleaned_matches:
        ats = safe_number(match.get("ats_score"), 0)
        verdict = clean_cell(match.get("verdict"))
        final_remark = clean_cell(match.get("final_remark"))

        is_relevant = (
            ats >= 50
            and verdict != "Not Suitable"
            and final_remark != "Not Suitable"
        )

        if is_relevant:
            relevant_matches.append(match)

    if relevant_matches:
        return relevant_matches

    # Nothing relevant: keep only closest/best one.
    best_one = cleaned_matches[0]
    best_one["verdict"] = "Not Suitable"
    best_one["final_remark"] = "Not Suitable"
    best_one["call_status"] = ""

    return [best_one]


# =========================================================
# STATUS HELPERS
# =========================================================

def get_default_status():
    return {
        "job_id": "",
        "pipeline": "hybrid",
        "status": "idle",  # idle, queued, processing, completed, failed
        "message": "No active hybrid job",
        "total": 0,
        "processed": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "current_file": "",
        "current_batch": "",
        "output_filename": "",
        "download_url": "",
        "failed_files": [],
        "skipped_files": [],
        "token_usage": {
            "input_tokens": 0,
            "output_tokens": 0,
            "total_tokens": 0,
        },
        "started_at": "",
        "updated_at": "",
        "completed_at": "",
        "cost_info": {
            "model_name": HYBRID_CLAUDE_MODEL,
            "input_cost_usd": 0,
            "output_cost_usd": 0,
            "total_cost_usd": 0,
            "total_cost_inr": 0,
            "cost_per_resume_usd": 0,
            "cost_per_resume_inr": 0,
        },
        "runtime_info": {
            "started_at": "",
            "completed_at": "",
            "total_seconds": 0,
            "total_time_text": "",
            "average_seconds_per_resume": 0,
        },
        "resume_logs": [],
        "recent_logs": [],
    }


def read_status():
    if not os.path.exists(HYBRID_STATUS_PATH):
        return get_default_status()

    try:
        with open(HYBRID_STATUS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return get_default_status()


def write_status(data: dict):
    data["updated_at"] = datetime.now().isoformat()

    with open(HYBRID_STATUS_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def add_recent_log(message: str):
    status = read_status()
    logs = status.get("recent_logs", [])
    logs.append({
        "time": datetime.now().strftime("%H:%M:%S"),
        "message": message,
    })
    status["recent_logs"] = logs[-40:]
    status["message"] = message
    write_status(status)
    debug(message)


def is_active_job_running():
    status = read_status()
    return status.get("status") in ["queued", "processing"]


def reset_hybrid_job_storage():
    if os.path.exists(HYBRID_STATUS_PATH):
        os.remove(HYBRID_STATUS_PATH)

    if os.path.exists(HYBRID_UPLOAD_DIR):
        shutil.rmtree(HYBRID_UPLOAD_DIR)

    os.makedirs(HYBRID_UPLOAD_DIR, exist_ok=True)


def format_duration(seconds: float):
    seconds = int(seconds)
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60

    if hours > 0:
        return f"{hours} hr {minutes} min {secs} sec"

    if minutes > 0:
        return f"{minutes} min {secs} sec"

    return f"{secs} sec"


def calculate_cost_info(
    input_tokens: int,
    output_tokens: int,
    processed_count: int,
    model_name: str = HYBRID_CLAUDE_MODEL,
):
    input_cost_usd = (input_tokens / 1_000_000) * HYBRID_INPUT_RATE_PER_MILLION_USD
    output_cost_usd = (output_tokens / 1_000_000) * HYBRID_OUTPUT_RATE_PER_MILLION_USD
    total_cost_usd = input_cost_usd + output_cost_usd
    total_cost_inr = total_cost_usd * USD_TO_INR

    cost_per_resume_usd = 0
    cost_per_resume_inr = 0

    if processed_count > 0:
        cost_per_resume_usd = total_cost_usd / processed_count
        cost_per_resume_inr = total_cost_inr / processed_count

    return {
        "model_name": model_name,
        "input_rate_per_million_usd": HYBRID_INPUT_RATE_PER_MILLION_USD,
        "output_rate_per_million_usd": HYBRID_OUTPUT_RATE_PER_MILLION_USD,
        "usd_to_inr": USD_TO_INR,
        "input_cost_usd": round(input_cost_usd, 4),
        "output_cost_usd": round(output_cost_usd, 4),
        "total_cost_usd": round(total_cost_usd, 4),
        "total_cost_inr": round(total_cost_inr, 2),
        "cost_per_resume_usd": round(cost_per_resume_usd, 4),
        "cost_per_resume_inr": round(cost_per_resume_inr, 2),
        "note": "Hybrid estimated cost based on configured Haiku pricing. Final provider billing may vary.",
    }


# =========================================================
# REQUIREMENT EXCEL LOADING
# =========================================================

def load_requirement_df(cfg: Settings) -> pd.DataFrame:
    excel_path = UPLOADED_REQUIREMENT_PATH

    debug(f"Loading uploaded requirement Excel: {excel_path}")

    if not os.path.exists(excel_path):
        raise HTTPException(
            status_code=400,
            detail="Requirement Excel not uploaded. Please upload Requirement Excel first from main app.",
        )

    xls = pd.ExcelFile(excel_path)
    debug(f"Available sheets: {xls.sheet_names}")

    matched_sheets = [
        s for s in xls.sheet_names if "requirement" in str(s).strip().lower()
    ]

    if matched_sheets:
        sheet_name = matched_sheets[0]
    else:
        sheet_name = xls.sheet_names[0]

    debug(f"Using requirement sheet: {sheet_name}")

    df = pd.read_excel(excel_path, sheet_name=sheet_name)
    df.columns = [str(c).strip() for c in df.columns]

    rename_map = {
        "Monthly Company Pay Rate": "Rate Card",
        "Monthly Company Pay Rate ": "Rate Card",
        "Anually Company Pay Rate": "Yearly Rate",
        "Annually Company Pay Rate": "Yearly Rate",
        "Annual Company Pay Rate": "Yearly Rate",
        "Request ID": "Request-ID",
        "Request Id": "Request-ID",
        "Name of Skills": "Skills - Name",
        "Job Position": "Job Title",
        "Name w/o Link": "MSP Owner",
        "Work Location City": "Work Location City",
        "Work Location CDF": "Work Location CDF",
    }

    df = df.rename(columns=rename_map)

    missing_cols = [
        col for col in REQUIRED_REQUIREMENT_COLUMNS if col not in df.columns
    ]

    # In hybrid, do not fail for columns we can keep blank after rename,
    # because CEO style sheets can have slightly different headers.
    for col in missing_cols:
        df[col] = ""

    df["Request-ID"] = df["Request-ID"].astype(str).str.strip()
    df = df[df["Request-ID"].astype(str).str.strip() != ""]
    df = df.dropna(how="all")

    optional_cols = [
        "Work Location City",
        "System Enhancements Required",
        "Candidate Annual CTC",
    ]

    for col in optional_cols:
        if col not in df.columns:
            df[col] = ""

    if df.empty:
        raise HTTPException(
            status_code=400,
            detail="Requirement Excel has no valid rows with Request-ID.",
        )

    debug(f"Requirement dataframe shape: {df.shape}")
    debug(f"Requirement columns: {list(df.columns)}")

    return df


def find_requirement_by_request_id(
    requirements_df: pd.DataFrame, request_id: str
) -> dict:
    request_id = clean_cell(request_id)

    if not request_id:
        return {}

    matched = requirements_df[
        requirements_df["Request-ID"].astype(str).str.strip() == request_id
    ]

    if matched.empty:
        return {}

    return matched.iloc[0].fillna("").to_dict()


# =========================================================
# LOCAL RESUME PARSING
# =========================================================

def extract_text_from_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()
    except Exception as e:
        debug(f"PDF parse failed for {path}: {repr(e)}")
        return ""


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document

        doc = Document(path)
        parts = []

        for p in doc.paragraphs:
            parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)

        return "\n".join(parts).strip()
    except Exception as e:
        debug(f"DOCX parse failed for {path}: {repr(e)}")
        return ""


def convert_office_file_to_docx_or_pdf_text(path: str) -> str:
    """
    For .doc files, try LibreOffice conversion to docx/pdf.
    This keeps hybrid more practical without changing main.py.
    """
    ext = Path(path).suffix.lower()

    if ext != ".doc":
        return ""

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            command = [
                "libreoffice",
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                path,
            ]

            result = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120,
            )

            if result.returncode != 0:
                debug(f"DOC to DOCX conversion failed: {result.stderr}")
                return ""

            docx_files = [f for f in os.listdir(temp_dir) if f.lower().endswith(".docx")]

            if not docx_files:
                return ""

            converted_path = os.path.join(temp_dir, docx_files[0])
            return extract_text_from_docx(converted_path)

    except Exception as e:
        debug(f"DOC conversion exception: {repr(e)}")
        return ""


def extract_resume_text(path: str) -> str:
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)

    if ext == ".docx":
        return extract_text_from_docx(path)

    if ext == ".doc":
        return convert_office_file_to_docx_or_pdf_text(path)

    return ""


# =========================================================
# LOCAL CANDIDATE EXTRACTION
# =========================================================

SKILL_VOCAB = [
    "aws", "azure", "gcp", "kubernetes", "aks", "eks", "gke", "docker", "terraform",
    "ansible", "jenkins", "github actions", "gitlab", "ci/cd", "helm", "prometheus",
    "grafana", "datadog", "splunk", "linux", "rhel", "ubuntu", "bash", "shell",
    "devops", "sre", "iac", "cloudformation",
    "python", "java", "javascript", "typescript", "c#", "c++", ".net", "golang",
    "scala", "ruby", "php", "swift", "kotlin", "rust", "perl",
    "react", "angular", "vue", "nextjs", "node", "node js", "html", "css", "spartacus",
    "fastapi", "flask", "django", "express",
    "sql", "oracle", "postgres", "postgresql", "mysql", "mongodb", "cassandra",
    "redis", "snowflake", "databricks", "spark", "hadoop", "hive", "kafka",
    "airflow", "dbt", "tableau", "power bi", "qlik", "informatica", "iics",
    "teradata", "etl", "data warehouse", "data lake", "bigquery",
    "sap", "s/4hana", "s4hana", "vistex", "hana", "abap", "sap mm", "sap sd",
    "sap fico", "sap crm", "sap bw", "datasphere", "sap hybris", "sap commerce",
    "ariba", "successfactors",
    "mdm", "master data", "data modeling", "erwin", "dimensional modeling",
    "machine learning", "ml", "tensorflow", "pytorch", "llm", "nlp", "genai",
    "generative ai", "scikit", "pandas", "numpy",
    "selenium", "cypress", "jmeter", "qa", "automation testing", "appium",
    "workday", "salesforce", "servicenow", "oracle fusion", "netsuite",
    "gxp", "gmp", "clinical", "cdisc", "sdtm", "rbqm", "etmf", "ich",
    "21 cfr part 11", "veeva", "pharma", "pharmacovigilance", "regulatory",
    "jira", "confluence", "git", "agile", "scrum", "itil", "kanban",
    "nexthink", "citrix", "vmware", "sccm", "active directory",
]


def extract_email(text: str) -> str:
    m = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else ""


def extract_phone(text: str) -> str:
    patterns = [
        r"(\+91[\s-]?\d{10})",
        r"(\+91[\s-]?\d{5}[\s-]?\d{5})",
        r"\b([6-9]\d{9})\b",
        r"(\+\d{1,3}[\s-]?\d{6,14})",
    ]

    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return m.group(1)

    return ""


def extract_years(text: str) -> Optional[float]:
    t = text.lower()

    patterns = [
        r"(\d{1,2}(?:\.\d+)?)\s*\+\s*years",
        r"(\d{1,2}(?:\.\d+)?)\s*years\s*of\s*experience",
        r"total\s*experience[:\s]*(\d{1,2}(?:\.\d+)?)",
        r"experience[:\s]*(\d{1,2}(?:\.\d+)?)\s*years",
        r"(\d{1,2}(?:\.\d+)?)\s*yrs",
    ]

    for pattern in patterns:
        m = re.search(pattern, t)
        if m:
            return safe_number(m.group(1), None)

    return None


def extract_name(text: str, filename: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    bad_words = (
        "resume", "curriculum", "email", "phone", "mobile", "contact",
        "linkedin", "github", "summary", "objective", "experience",
        "skills", "education", "profile", "candidate", "address",
    )

    for line in lines[:20]:
        low = line.lower()

        if any(low.startswith(b) or b in low[:25] for b in bad_words):
            continue

        if "@" in line or re.search(r"\d", line):
            continue

        words = line.split()

        if 2 <= len(words) <= 4 and len(line) <= 60:
            return line

    return Path(filename).stem


def extract_location(text: str) -> dict:
    """
    Hybrid local location extraction.
    Conservative: avoid using random employer/project locations when possible.
    """
    cities = [
        "Mumbai", "Navi Mumbai", "Thane", "Pune", "Bangalore", "Bengaluru",
        "Hyderabad", "Chennai", "Delhi", "New Delhi", "Noida", "Gurgaon", "Gurugram",
        "Kolkata", "Ahmedabad", "Vadodara", "Indore", "Nagpur", "Kochi",
        "Coimbatore", "Jaipur", "Lucknow", "Bhubaneswar", "Mysore", "Mysuru",
    ]

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    top_text = "\n".join(lines[:25])
    top_lower = top_text.lower()

    explicit_patterns = [
        r"current location[:\s]+([a-zA-Z\s,.-]+)",
        r"location[:\s]+([a-zA-Z\s,.-]+)",
        r"address[:\s]+([a-zA-Z\s,.-]+)",
        r"residence[:\s]+([a-zA-Z\s,.-]+)",
        r"preferred location[:\s]+([a-zA-Z\s,.-]+)",
    ]

    for pattern in explicit_patterns:
        m = re.search(pattern, top_lower)
        if m:
            value = m.group(1)[:100]
            for city in cities:
                if city.lower() in value:
                    return {
                        "candidate_location": city,
                        "candidate_city": city,
                        "candidate_state": None,
                        "candidate_country": "India",
                        "candidate_location_source": "local_text_extraction",
                        "candidate_location_confidence": "medium",
                    }

    # fallback only from top/header area
    for city in cities:
        if re.search(rf"\b{re.escape(city.lower())}\b", top_lower):
            return {
                "candidate_location": city,
                "candidate_city": city,
                "candidate_state": None,
                "candidate_country": "India",
                "candidate_location_source": "local_text_extraction_header",
                "candidate_location_confidence": "medium",
            }

    return {
        "candidate_location": None,
        "candidate_city": None,
        "candidate_state": None,
        "candidate_country": None,
        "candidate_location_source": "not mentioned",
        "candidate_location_confidence": "not_found",
    }


def extract_skills(text: str) -> List[str]:
    text_norm = normalize_text(text)
    found = set()

    for skill in SKILL_VOCAB:
        skill_norm = normalize_text(skill)

        if not skill_norm:
            continue

        if any(c in skill_norm for c in (" ", ".", "/", "+", "#")):
            if skill_norm in text_norm:
                found.add(skill)
        else:
            if re.search(rf"\b{re.escape(skill_norm)}\b", text_norm):
                found.add(skill)

    return sorted(found)


def extract_recent_job_title(text: str) -> str:
    # Cheap heuristic only; Claude matching will reason with profile_summary.
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title_keywords = [
        "developer", "engineer", "consultant", "analyst", "manager",
        "architect", "lead", "specialist", "administrator", "tester",
        "scientist", "data", "devops", "sap", "qa", "automation",
    ]

    for line in lines[:80]:
        low = line.lower()
        if any(k in low for k in title_keywords) and len(line) <= 90:
            return line

    return ""


def build_candidate_info_from_text(text: str, filename: str) -> dict:
    clean_text = re.sub(r"\s+", " ", text).strip()
    loc = extract_location(text)

    candidate_info = {
        "candidate_name": extract_name(text, filename),
        "candidate_phone": extract_phone(clean_text),
        "candidate_email": extract_email(clean_text),

        "candidate_location": loc.get("candidate_location"),
        "candidate_city": loc.get("candidate_city"),
        "candidate_state": loc.get("candidate_state"),
        "candidate_country": loc.get("candidate_country"),
        "candidate_location_source": loc.get("candidate_location_source"),
        "candidate_location_evidence": loc.get("candidate_location"),
        "candidate_location_confidence": loc.get("candidate_location_confidence"),

        "candidate_total_experience_years": extract_years(clean_text),
        "candidate_current_ctc": None,
        "candidate_expected_ctc": None,
        "candidate_notice_period": None,

        "candidate_skills": extract_skills(clean_text),
        "recent_job_title": extract_recent_job_title(text),
        "profile_summary": clean_text[:3500],
        "local_parse_used": True,
    }

    return candidate_info


# =========================================================
# PYTHON SHORTLISTING - SAME PURPOSE AS MAIN APP
# =========================================================

def build_requirement_search_text(row: dict) -> str:
    return " ".join(
        [
            clean_cell(row.get("Job Title")),
            clean_cell(row.get("Skills - Name")),
            clean_cell(row.get("Skills - Experience")),
            clean_cell(row.get("Additional Skills")),
            clean_cell(row.get("Job Description")),
            clean_cell(row.get("Work Location CDF")),
        ]
    )


def build_candidate_search_text(candidate_info: dict) -> str:
    return " ".join(
        [
            " ".join(candidate_info.get("candidate_skills") or []),
            clean_cell(candidate_info.get("recent_job_title")),
            clean_cell(candidate_info.get("profile_summary")),
            clean_cell(candidate_info.get("candidate_location")),
        ]
    )


def shortlist_requirements(
    requirements_df: pd.DataFrame,
    candidate_info: dict,
    top_n: int = HYBRID_TOP_N_REQUIREMENTS,
) -> list:
    candidate_text = build_candidate_search_text(candidate_info)
    candidate_norm = normalize_text(candidate_text)
    candidate_tokens = set(candidate_norm.split())

    scored_rows = []

    for _, row in requirements_df.iterrows():
        row_dict = row.fillna("").to_dict()

        req_text = build_requirement_search_text(row_dict)
        req_norm = normalize_text(req_text)
        req_tokens = set(req_norm.split())

        overlap = candidate_tokens.intersection(req_tokens)
        score = len(overlap)

        # Strong boost on exact primary skill tokens
        primary_skills = split_skills(row_dict.get("Skills - Name"))
        for skill in primary_skills:
            skill_norm = normalize_text(skill)
            if skill_norm and skill_norm in candidate_norm:
                score += 25

        # Boost additional skills
        additional_skills = split_skills(row_dict.get("Additional Skills"))
        for skill in additional_skills[:80]:
            skill_norm = normalize_text(skill)
            if skill_norm and skill_norm in candidate_norm:
                score += 5

        # Prefer open requirements slightly
        status = normalize_text(row_dict.get("Status"))
        if status == "open":
            score += 5

        # Location small boost only if candidate location exists
        candidate_city = normalize_text(candidate_info.get("candidate_city"))
        req_location = normalize_text(row_dict.get("Work Location CDF"))
        if candidate_city and candidate_city in req_location:
            score += 5

        scored_rows.append((score, row_dict))

    scored_rows = sorted(scored_rows, key=lambda x: x[0], reverse=True)

    top_rows = [row for score, row in scored_rows[:top_n]]

    debug(f"Hybrid shortlisted {len(top_rows)} requirements")
    return top_rows


# =========================================================
# CLAUDE CLIENT + HYBRID MATCHING
# =========================================================

def get_claude_client(cfg: Settings) -> anthropic.Anthropic:
    return anthropic.Anthropic(api_key=cfg.anthropic_api_key)


def requirement_for_prompt(row: dict) -> dict:
    return {
        "Request-ID": clean_cell(row.get("Request-ID")),
        "MSP Owner": clean_cell(row.get("MSP Owner")),
        "Job Title": clean_cell(row.get("Job Title")),
        "Skills - Name": clean_cell(row.get("Skills - Name")),
        "Skills - Experience": clean_cell(row.get("Skills - Experience")),
        "Additional Skills": clean_cell(row.get("Additional Skills")),
        "Job Description": clean_cell(row.get("Job Description"))[:3500],
        "Status": clean_cell(row.get("Status")),
        "Work Location City": clean_cell(row.get("Work Location City")),
        "Work Location CDF": clean_cell(row.get("Work Location CDF")),
        "Rate Card": clean_cell(row.get("Rate Card")),
        "Yearly Rate": clean_cell(row.get("Yearly Rate")),
        "System Enhancements Required": clean_cell(row.get("System Enhancements Required")),
        "Candidate Annual CTC": clean_cell(row.get("Candidate Annual CTC")),
    }


def call_claude_relevant_requirement_matches_hybrid(
    candidate_info: dict,
    shortlisted_requirements: list,
    cfg: Settings,
) -> dict:
    """
    Hybrid equivalent of main app's call_claude_relevant_requirement_matches.

    Difference:
    - Candidate profile is locally extracted / compressed.
    - Requirements are shortlisted locally.
    - Uses cheaper Claude model.
    - Still asks Claude to evaluate EVERY shortlisted requirement.
    - Python post-processing enforces all relevant else best one.
    """

    debug("Hybrid Claude relevant requirement matching started")

    client = get_claude_client(cfg)

    requirements_for_prompt = [
        requirement_for_prompt(row) for row in shortlisted_requirements
    ]

    prompt = f"""
You are an expert technical recruiter.

You are given:
1. One candidate profile extracted from resume using local parsing.
2. A shortlisted list of job requirements from Requirement Sheet.

Your task:
- Evaluate the candidate against EVERY shortlisted Request-ID/job requirement.
- Return match objects for all requirements you evaluated.
- A candidate can be mapped to multiple requirements when relevant.
- If no requirement is a good fit, still return the closest/best one as Not Suitable.
- Generate ATS score from 0 to 100 for each returned requirement.
- Generate recruiter-style verdict, call_status and remark for each returned requirement.

Critical business rule:
- Relevant matches normally have ATS >= 50 and verdict Possible Fit / Good Fit / Strong Fit.
- If multiple requirements are relevant, return all relevant ones.
- If nothing is relevant, return only the closest/best one as Not Suitable.
- Do not invent Request-ID.
- Use ONLY provided requirements.

Important:
- NPU means Not Picked Up. It cannot be detected from resume. Do not set NPU unless input clearly says call status is NPU.
- Use High CTC only if candidate CTC is available and exceeds requirement budget.
- Use High Notice Period only if notice period is available and clearly high.
- Use Location Mismatch if location is main issue.
- Use Skill mismatch if skill gap is main issue.
- Use Exp mismatch if experience is main issue.
- Use Match only when candidate is a strong/clear fit.
- Use Not Suitable when profile/domain is not aligned.
- call_status must always be empty string "" because call status is recruiter-entered manually after calling the candidate.

ATS scoring guidance:
- ATS must reflect overall match quality between the candidate and each selected requirement.
- Do not give same ATS score to every requirement unless their fit is genuinely almost identical.
- Give lower score when profile/domain is weakly aligned.
- Give medium score when candidate is partially aligned but has important gaps.
- Give higher score when candidate strongly matches experience, core skills, role focus, and location/availability where available.
- Keep ATS realistic and recruiter-friendly.

Location rules:
- Compare candidate location with requirement location only when candidate_location_confidence is "high" or "medium".
- Do not compare location if candidate_location_confidence is "low" or "not_found".
- If candidate_location is null/blank/not_found, set location_mismatch = "Not Evaluated".
- If location is not evaluated, do not use final_remark = "Location Mismatch".
- Compare locations semantically and case-insensitively.
- Ignore formatting differences, symbols, punctuation, country/state suffixes, office names, SEZ/branch/building names, and extra whitespace.
- Treat equivalent city spellings or common regional names as same when clearly equivalent.
- Do not infer candidate location from employer/work/education/project location.
- In reason, if candidate location is missing, write:
  "Candidate location is not explicitly mentioned in the resume, so location fit is not evaluated."

Allowed verdict values:
- Strong Fit
- Good Fit
- Possible Fit
- Not Suitable

Allowed final_remark values:
- Exp mismatch
- High CTC
- High Notice Period
- Location Mismatch
- Match
- Not Suitable
- Skill mismatch

Set experience_mismatch as "Yes" if candidate total experience is less than required experience, otherwise "No".
Set skill_mismatch as "Yes" if candidate is missing important required skills, otherwise "No".
Set location_mismatch as "Yes" only if candidate location is explicitly available and conflicts with requirement location.
If candidate location is missing, set location_mismatch as "Not Evaluated".

Candidate Profile:
{json.dumps(candidate_info, ensure_ascii=False, indent=2)}

Shortlisted Requirements:
{json.dumps(requirements_for_prompt, ensure_ascii=False, indent=2)}

Return only valid JSON:
{{
  "matches": [
    {{
      "request_id": null,
      "ats_score": 0,
      "verdict": "Not Suitable",
      "call_status": "",
      "final_remark": "Not Suitable",
      "experience_mismatch": "No",
      "skill_mismatch": "No",
      "location_mismatch": "Not Evaluated",
      "matching_skills": [],
      "missing_skills": [],
      "reason": ""
    }}
  ]
}}
"""

    message = client.messages.create(
        model=HYBRID_CLAUDE_MODEL,
        max_tokens=HYBRID_MAX_TOKENS,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt,
                    }
                ],
            }
        ],
    )

    raw = message.content[0].text.strip()

    usage = getattr(message, "usage", None)
    input_tokens = getattr(usage, "input_tokens", 0) if usage else 0
    output_tokens = getattr(usage, "output_tokens", 0) if usage else 0

    try:
        result = clean_json_text(raw)
    except Exception as e:
        debug(f"Hybrid multi-match JSON parse failed: {repr(e)}")
        result = {"matches": []}

    matches = result.get("matches")
    if not isinstance(matches, list):
        matches = []

    cleaned_matches = []

    for match in matches:
        if not isinstance(match, dict):
            continue

        request_id = clean_cell(match.get("request_id")) or clean_cell(match.get("best_request_id"))
        if not request_id:
            continue

        match["request_id"] = request_id
        match["best_request_id"] = request_id
        match["call_status"] = ""

        final_remark = clean_cell(match.get("final_remark"))
        if final_remark not in VALID_REMARKS:
            final_remark = "Not Suitable"
        match["final_remark"] = final_remark

        ats = clamp_int(match.get("ats_score"), 0, 100)
        match["ats_score"] = ats

        # Keep verdict aligned with ATS, like main app
        if ats >= 85:
            verdict = "Strong Fit"
        elif ats >= 70:
            verdict = "Good Fit"
        elif ats >= 50:
            verdict = "Possible Fit"
        else:
            verdict = "Not Suitable"

        match["verdict"] = verdict

        # Safety: if candidate location is missing, never allow Location Mismatch.
        candidate_city = clean_cell(candidate_info.get("candidate_city"))
        candidate_location = clean_cell(candidate_info.get("candidate_location"))

        if not candidate_city and not candidate_location:
            if match.get("final_remark") == "Location Mismatch":
                match["final_remark"] = "Not Suitable"

            match["location_mismatch"] = "Not Evaluated"

            reason = clean_cell(match.get("reason"))
            if "location fit is not evaluated" not in reason.lower():
                match["reason"] = (
                    reason
                    + " Candidate location is not explicitly mentioned in the resume, so location fit is not evaluated."
                ).strip()

        if not clean_cell(match.get("experience_mismatch")):
            match["experience_mismatch"] = "No"

        if not clean_cell(match.get("skill_mismatch")):
            match["skill_mismatch"] = "No"

        if not clean_cell(match.get("location_mismatch")):
            match["location_mismatch"] = "Not Evaluated"

        cleaned_matches.append(match)

    # If Claude returned nothing, fallback to best shortlisted requirement.
    if not cleaned_matches and shortlisted_requirements:
        best_req = shortlisted_requirements[0]
        cleaned_matches = [
            {
                "request_id": clean_cell(best_req.get("Request-ID")),
                "best_request_id": clean_cell(best_req.get("Request-ID")),
                "ats_score": 0,
                "verdict": "Not Suitable",
                "call_status": "",
                "final_remark": "Not Suitable",
                "experience_mismatch": "No",
                "skill_mismatch": "Yes",
                "location_mismatch": "Not Evaluated",
                "matching_skills": [],
                "missing_skills": [],
                "reason": "No strong relevant match found in hybrid screening.",
            }
        ]

    # IMPORTANT: enforce same main app mapping rule here.
    final_matches = select_relevant_matches_or_best_one(cleaned_matches)

    debug(
        f"Hybrid Claude relevant requirement matching completed. "
        f"Raw matches={len(cleaned_matches)}, Final matches={len(final_matches)}"
    )

    return {
        "matches": final_matches,
        "_token_usage": {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": input_tokens + output_tokens,
        },
    }


# =========================================================
# OUTPUT ROW BUILDERS - SAME FASHION AS MAIN APP
# =========================================================

def build_output_row(
    req: dict,
    candidate_info: dict,
    match_result: dict,
    filename: str = "",
) -> dict:
    candidate_skills = candidate_info.get("candidate_skills") or []

    candidate_city = clean_cell(candidate_info.get("candidate_city"))
    candidate_location = clean_cell(candidate_info.get("candidate_location"))

    final_candidate_location = candidate_city or candidate_location

    return {
        "Request-ID": clean_cell(req.get("Request-ID")),
        "MSP Owner": clean_cell(req.get("MSP Owner")),
        "Job Title": clean_cell(req.get("Job Title")),
        "Skills - Name": clean_cell(req.get("Skills - Name")),
        "Skills - Experience": clean_cell(req.get("Skills - Experience")),
        "Additional Skills": clean_cell(req.get("Additional Skills")),
        "Job Description": clean_cell(req.get("Job Description")),
        "Work Location CDF": clean_cell(req.get("Work Location CDF")),
        "Rate Card": clean_cell(req.get("Rate Card")),
        "Annually": clean_cell(req.get("Yearly Rate")),
        "CV File Name": clean_cell(filename),
        "Candidate Name": clean_cell(candidate_info.get("candidate_name")),
        "Candidate Phone": clean_cell(candidate_info.get("candidate_phone")),
        "Candidate Email": clean_cell(candidate_info.get("candidate_email")),
        "Candidate Location": final_candidate_location,
        "Candidate Total Experience": clean_cell(candidate_info.get("candidate_total_experience_years")),
        "Candidate Skills": ", ".join(candidate_skills),
        "Experience Mismatch": clean_cell(match_result.get("experience_mismatch")) or "No",
        "Skill Mismatch": clean_cell(match_result.get("skill_mismatch")) or "No",
        "ATS": clean_cell(match_result.get("ats_score")),
        # Same as main.py reference: Output Sheet Remark contains recruiter reason.
        "Remark": clean_cell(match_result.get("reason")),
    }


def build_tracker_row(
    req: dict,
    candidate_info: dict,
    match_result: dict,
    filename: str,
) -> dict:
    candidate_name = clean_cell(candidate_info.get("candidate_name"))

    if not candidate_name:
        candidate_name = filename

    return {
        "Date": datetime.now().strftime("%d-%m-%Y"),
        "Request ID": clean_cell(req.get("Request-ID")) or clean_cell(match_result.get("best_request_id")),
        "Status": clean_cell(req.get("Status")),
        "Skills": clean_cell(req.get("Skills - Name")),
        "Candidate": candidate_name,
        "Verdict": clean_cell(match_result.get("verdict")),
        "Call Status": "",
        "Remarks": clean_cell(match_result.get("final_remark")) or "Not Suitable",
    }


# =========================================================
# EXCEL GENERATION - SAME SHEETS AS MAIN APP
# =========================================================

def style_output_sheet(ws):
    header_blue = PatternFill("solid", fgColor="B7DEE8")
    candidate_green = PatternFill("solid", fgColor="DAF2D0")
    result_pink = PatternFill("solid", fgColor="F2CEEF")
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.font = Font(bold=True, color="000000")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

        if cell.column <= 10:
            cell.fill = header_blue
        elif 11 <= cell.column <= 16:
            cell.fill = candidate_green
        else:
            cell.fill = result_pink

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def style_tracker_sheet(ws):
    header_fill = PatternFill("solid", fgColor="1F4E78")
    thin = Side(style="thin", color="808080")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    verdict_col = None
    for cell in ws[1]:
        if cell.value == "Verdict":
            verdict_col = cell.column
            break

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border

        if verdict_col:
            verdict_cell = row[verdict_col - 1]
            verdict = clean_cell(verdict_cell.value)
            color = VERDICT_COLORS.get(verdict)

            if color:
                verdict_cell.fill = PatternFill("solid", fgColor=color)
                verdict_cell.font = Font(bold=True, color="FFFFFF")

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def style_pivot_sheet(ws):
    header_fill = PatternFill("solid", fgColor="B7DEE8")
    total_fill = PatternFill("solid", fgColor="B7DEE8")
    thin = Side(style="thin", color="808080")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.font = Font(bold=True, color="000000")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="center")

        first_cell = row[0]
        if clean_cell(first_cell.value) == "Grand Total":
            for cell in row:
                cell.font = Font(bold=True)
                cell.fill = total_fill

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def auto_adjust_columns(ws):
    for col_idx, col_cells in enumerate(ws.columns, start=1):
        max_len = 0

        for cell in col_cells:
            value = clean_cell(cell.value)
            if len(value) > max_len:
                max_len = len(value[:100])

        width = min(max(max_len + 2, 12), 55)
        ws.column_dimensions[get_column_letter(col_idx)].width = width


def create_pivot_df(tracker_df: pd.DataFrame) -> pd.DataFrame:
    if tracker_df.empty or "Remarks" not in tracker_df.columns:
        rows = [
            {
                "Row Labels": "Grand Total",
                "Count of Remarks": 0,
            }
        ]
        return pd.DataFrame(rows, columns=["Row Labels", "Count of Remarks"])

    counts = tracker_df["Remarks"].fillna("Not Suitable").value_counts().to_dict()

    rows = []

    for remark in PIVOT_REMARK_ORDER:
        count = int(counts.get(remark, 0))
        if count > 0:
            rows.append(
                {
                    "Row Labels": remark,
                    "Count of Remarks": count,
                }
            )

    # Add unexpected remarks if any
    for remark, count in counts.items():
        if remark not in PIVOT_REMARK_ORDER:
            rows.append(
                {
                    "Row Labels": remark,
                    "Count of Remarks": int(count),
                }
            )

    grand_total = sum(row["Count of Remarks"] for row in rows)

    rows.append(
        {
            "Row Labels": "Grand Total",
            "Count of Remarks": grand_total,
        }
    )

    return pd.DataFrame(rows, columns=["Row Labels", "Count of Remarks"])


def create_final_excel(output_rows: list, tracker_rows: list, output_path: str):
    output_df = pd.DataFrame(output_rows, columns=OUTPUT_COLUMNS)
    tracker_df = pd.DataFrame(tracker_rows, columns=TRACKER_COLUMNS)
    pivot_df = create_pivot_df(tracker_df)

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        output_df.to_excel(writer, sheet_name="Output Sheet", index=False)
        tracker_df.to_excel(writer, sheet_name="Tracker", index=False)
        pivot_df.to_excel(writer, sheet_name="Pivot", index=False)

        workbook = writer.book

        ws_output = workbook["Output Sheet"]
        ws_tracker = workbook["Tracker"]
        ws_pivot = workbook["Pivot"]

        style_output_sheet(ws_output)
        style_tracker_sheet(ws_tracker)
        style_pivot_sheet(ws_pivot)

        auto_adjust_columns(ws_output)
        auto_adjust_columns(ws_tracker)
        auto_adjust_columns(ws_pivot)

    debug(f"Hybrid final Excel created: {output_path}")


# =========================================================
# FILE WRAPPER
# =========================================================

class SavedUploadFile:
    def __init__(self, filename: str, content_type: str, path: str):
        self.filename = filename
        self.content_type = content_type
        self.path = path

    async def read(self):
        with open(self.path, "rb") as f:
            return f.read()


# =========================================================
# SINGLE RESUME PROCESSING
# =========================================================

async def process_single_resume_file_hybrid(
    file: SavedUploadFile,
    requirements_df: pd.DataFrame,
    cfg: Settings,
    index: int,
    total_files: int,
) -> dict:
    filename = file.filename
    debug(f"[{index}/{total_files}] Hybrid processing started: {filename}")

    ext = get_file_extension(filename)

    if ext not in ALLOWED_RESUME_EXTENSIONS:
        return {
            "status": "skipped",
            "filename": filename,
            "reason": f"Unsupported file type: {ext}",
            "output_rows": [],
            "tracker_rows": [],
            "token_usage": {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0},
            "matched_requirements_count": 0,
        }

    text = extract_resume_text(file.path)

    if not text or len(text.strip()) < 100:
        return {
            "status": "skipped",
            "filename": filename,
            "reason": "Resume text extraction failed or text is too small.",
            "output_rows": [],
            "tracker_rows": [],
            "token_usage": {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0},
            "matched_requirements_count": 0,
        }

    candidate_info = build_candidate_info_from_text(text, filename)

    shortlisted_requirements = shortlist_requirements(
        requirements_df=requirements_df,
        candidate_info=candidate_info,
        top_n=HYBRID_TOP_N_REQUIREMENTS,
    )

    if not shortlisted_requirements:
        return {
            "status": "skipped",
            "filename": filename,
            "reason": "No requirements available for matching.",
            "output_rows": [],
            "tracker_rows": [],
            "token_usage": {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0},
            "matched_requirements_count": 0,
        }

    match_result = call_claude_relevant_requirement_matches_hybrid(
        candidate_info=candidate_info,
        shortlisted_requirements=shortlisted_requirements,
        cfg=cfg,
    )

    matches = match_result.get("matches", [])
    output_rows = []
    tracker_rows = []

    for match in matches:
        request_id = clean_cell(match.get("request_id")) or clean_cell(match.get("best_request_id"))
        req = find_requirement_by_request_id(requirements_df, request_id)

        if not req:
            debug(f"Requirement not found for request_id={request_id}")
            continue

        output_rows.append(
            build_output_row(
                req=req,
                candidate_info=candidate_info,
                match_result=match,
                filename=filename,
            )
        )

        tracker_rows.append(
            build_tracker_row(
                req=req,
                candidate_info=candidate_info,
                match_result=match,
                filename=filename,
            )
        )

    debug(
        f"[{index}/{total_files}] Hybrid processing completed: {filename}, "
        f"matched_rows={len(output_rows)}"
    )

    return {
        "status": "processed",
        "filename": filename,
        "output_rows": output_rows,
        "tracker_rows": tracker_rows,
        "token_usage": match_result.get("_token_usage", {}),
        "matched_requirements_count": len(output_rows),
    }


# =========================================================
# BACKGROUND JOB
# =========================================================

async def run_hybrid_background_job(
    job_id: str,
    saved_files: list,
    cfg: Settings,
):
    start_time = time.time()

    output_rows = []
    tracker_rows = []

    total_input_tokens = 0
    total_output_tokens = 0
    total_tokens = 0

    failed_files = []
    skipped_files = []

    try:
        total_files = len(saved_files)

        status = read_status()
        status.update(
            {
                "status": "processing",
                "message": "Hybrid processing resumes",
                "processed": 0,
                "successful": 0,
                "failed": 0,
                "skipped": 0,
                "current_file": "",
                "current_batch": f"0/{total_files}",
                "resume_logs": [],
                "runtime_info": {
                    "started_at": status.get("started_at", datetime.now().isoformat()),
                    "completed_at": "",
                    "total_seconds": 0,
                    "total_time_text": "",
                    "average_seconds_per_resume": 0,
                },
                "cost_info": calculate_cost_info(
                    input_tokens=0,
                    output_tokens=0,
                    processed_count=0,
                    model_name=HYBRID_CLAUDE_MODEL,
                ),
            }
        )
        write_status(status)

        add_recent_log("Loading requirement Excel for hybrid pipeline")
        requirements_df = load_requirement_df(cfg)

        semaphore = asyncio.Semaphore(HYBRID_MAX_PARALLEL)

        async def process_with_limit(index: int, item: dict):
            async with semaphore:
                filename = item["filename"]

                fake_file = SavedUploadFile(
                    filename=filename,
                    content_type=item.get("content_type") or "application/pdf",
                    path=item["path"],
                )

                return await process_single_resume_file_hybrid(
                    file=fake_file,
                    requirements_df=requirements_df,
                    cfg=cfg,
                    index=index,
                    total_files=total_files,
                )

        # Sequential by default because HYBRID_MAX_PARALLEL=1.
        for index, item in enumerate(saved_files, start=1):
            filename = item["filename"]
            resume_start_time = time.time()
            resume_started_at = datetime.now().isoformat()

            status = read_status()
            status.update(
                {
                    "status": "processing",
                    "message": f"Hybrid processing resume {index} of {total_files}",
                    "current_file": filename,
                    "current_batch": f"{index}/{total_files}",
                }
            )
            write_status(status)

            try:
                result = await process_with_limit(index, item)
            except Exception as e:
                debug(f"Hybrid exception for {filename}: {repr(e)}")
                result = {
                    "status": "failed",
                    "filename": filename,
                    "reason": str(e),
                    "output_rows": [],
                    "tracker_rows": [],
                    "token_usage": {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0},
                    "matched_requirements_count": 0,
                }

            resume_end_time = time.time()
            resume_seconds = round(resume_end_time - resume_start_time, 2)

            usage = result.get("token_usage", {}) or {}

            resume_input_tokens = int(usage.get("input_tokens", 0) or 0)
            resume_output_tokens = int(usage.get("output_tokens", 0) or 0)
            resume_total_tokens = int(usage.get("total_tokens", 0) or 0)

            total_input_tokens += resume_input_tokens
            total_output_tokens += resume_output_tokens
            total_tokens += resume_total_tokens

            output_rows.extend(result.get("output_rows", []))
            tracker_rows.extend(result.get("tracker_rows", []))

            if result.get("status") == "skipped":
                skipped_files.append(result.get("filename", filename))
            elif result.get("status") == "failed":
                failed_files.append(result.get("filename", filename))
            else:
                # processed means resume was attempted successfully
                pass

            resume_log = {
                "index": index,
                "filename": result.get("filename", filename),
                "status": result.get("status", "processed"),
                "started_at": resume_started_at,
                "completed_at": datetime.now().isoformat(),
                "duration_seconds": resume_seconds,
                "duration_text": format_duration(resume_seconds),
                "input_tokens": resume_input_tokens,
                "output_tokens": resume_output_tokens,
                "total_tokens": resume_total_tokens,
                "matched_requirements_count": result.get("matched_requirements_count", 0),
                "reason": result.get("reason", ""),
            }

            successful_count = index - len(skipped_files) - len(failed_files)
            processed_count_for_cost = max(index - len(skipped_files) - len(failed_files), 1)

            status = read_status()
            resume_logs = status.get("resume_logs", [])
            resume_logs.append(resume_log)

            cost_info = calculate_cost_info(
                input_tokens=total_input_tokens,
                output_tokens=total_output_tokens,
                processed_count=processed_count_for_cost,
                model_name=HYBRID_CLAUDE_MODEL,
            )

            status.update(
                {
                    "status": "processing",
                    "message": f"Hybrid processed resume {index} of {total_files}",
                    "processed": index,
                    "successful": successful_count,
                    "failed": len(failed_files),
                    "skipped": len(skipped_files),
                    "current_file": filename,
                    "current_batch": f"{index}/{total_files}",
                    "failed_files": failed_files,
                    "skipped_files": skipped_files,
                    "token_usage": {
                        "input_tokens": total_input_tokens,
                        "output_tokens": total_output_tokens,
                        "total_tokens": total_tokens,
                    },
                    "cost_info": cost_info,
                    # For frontend compatibility with earlier hybrid page
                    "token_usage_flat": {
                        "input_tokens": total_input_tokens,
                        "output_tokens": total_output_tokens,
                        "total_tokens": total_tokens,
                        "cost_usd": cost_info["total_cost_usd"],
                        "cost_inr": cost_info["total_cost_inr"],
                        "cost_per_resume_inr": cost_info["cost_per_resume_inr"],
                    },
                    "resume_logs": resume_logs,
                }
            )

            # Also include cost fields inside token_usage because some frontend versions expect this.
            status["token_usage"] = {
                "input_tokens": total_input_tokens,
                "output_tokens": total_output_tokens,
                "total_tokens": total_tokens,
                "cost_usd": cost_info["total_cost_usd"],
                "cost_inr": cost_info["total_cost_inr"],
                "cost_per_resume_inr": cost_info["cost_per_resume_inr"],
            }

            write_status(status)

        if not output_rows:
            raise RuntimeError("Hybrid completed but no output rows were generated.")

        output_filename = f"{POLLING_OUTPUT_PREFIX}_{job_id}.xlsx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        create_final_excel(output_rows, tracker_rows, output_path)

        total_seconds = round(time.time() - start_time, 2)
        successful_count = total_files - len(skipped_files) - len(failed_files)
        avg_seconds = round(total_seconds / max(total_files, 1), 2)

        final_cost_info = calculate_cost_info(
            input_tokens=total_input_tokens,
            output_tokens=total_output_tokens,
            processed_count=max(successful_count, 1),
            model_name=HYBRID_CLAUDE_MODEL,
        )

        completed_at = datetime.now().isoformat()

        status = read_status()
        status.update(
            {
                "status": "completed",
                "message": "Hybrid resume analysis completed successfully.",
                "processed": total_files,
                "successful": successful_count,
                "failed": len(failed_files),
                "skipped": len(skipped_files),
                "current_file": "",
                "current_batch": f"{total_files}/{total_files}",
                "output_filename": output_filename,
                "download_url": f"/hybrid-download/{output_filename}",
                "failed_files": failed_files,
                "skipped_files": skipped_files,
                "token_usage": {
                    "input_tokens": total_input_tokens,
                    "output_tokens": total_output_tokens,
                    "total_tokens": total_tokens,
                    "cost_usd": final_cost_info["total_cost_usd"],
                    "cost_inr": final_cost_info["total_cost_inr"],
                    "cost_per_resume_inr": final_cost_info["cost_per_resume_inr"],
                },
                "cost_info": final_cost_info,
                "runtime_info": {
                    "started_at": status.get("started_at", ""),
                    "completed_at": completed_at,
                    "total_seconds": total_seconds,
                    "total_time_text": format_duration(total_seconds),
                    "average_seconds_per_resume": avg_seconds,
                },
                "completed_at": completed_at,
                "processing_time_seconds": total_seconds,
            }
        )
        write_status(status)

        add_recent_log("Hybrid resume analysis completed successfully")

    except Exception as e:
        debug(f"Hybrid background job failed: {repr(e)}")
        total_seconds = round(time.time() - start_time, 2)

        status = read_status()
        status.update(
            {
                "status": "failed",
                "message": f"Hybrid resume analysis failed: {str(e)}",
                "completed_at": datetime.now().isoformat(),
                "runtime_info": {
                    "started_at": status.get("started_at", ""),
                    "completed_at": datetime.now().isoformat(),
                    "total_seconds": total_seconds,
                    "total_time_text": format_duration(total_seconds),
                    "average_seconds_per_resume": 0,
                },
            }
        )
        write_status(status)


# =========================================================
# ROUTES
# =========================================================

@app.get("/")
def root():
    return {
        "message": "Hybrid Resume Requirement Matching API",
        "usage": {
            "health": "GET /hybrid-health",
            "start": "POST /hybrid-start-bulk-analyze with resume files",
            "status": "GET /hybrid-status",
            "download": "GET /hybrid-download/{filename}",
        },
    }


@app.get("/hybrid-health")
def hybrid_health():
    return {
        "status": "ok",
        "pipeline": "hybrid",
        "model": HYBRID_CLAUDE_MODEL,
        "top_n_requirements": HYBRID_TOP_N_REQUIREMENTS,
        "max_parallel": HYBRID_MAX_PARALLEL,
        "business_rule": "all relevant matches if ATS >= 50 else one closest/best fit",
        "excel_sheets": ["Output Sheet", "Tracker", "Pivot"],
    }


@app.post("/hybrid-start-bulk-analyze", response_model=HybridStartResponse)
async def hybrid_start_bulk_analyze(
    files: List[UploadFile] = File(...),
    cfg: Settings = Depends(get_settings),
):
    debug("=" * 100)
    debug("Hybrid start bulk analyze request received")
    debug(f"Total files received: {len(files)}")

    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")

    if not os.path.exists(UPLOADED_REQUIREMENT_PATH):
        raise HTTPException(
            status_code=400,
            detail="Please upload valid Requirement Excel from main app before running hybrid analysis.",
        )

    if is_active_job_running():
        raise HTTPException(
            status_code=409,
            detail="A hybrid resume analysis job is already running. Please wait until it completes.",
        )

    reset_hybrid_job_storage()

    job_id = f"HYBRID_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{str(uuid.uuid4())[:8]}"

    saved_files = []
    skipped_files = []

    for index, file in enumerate(files, start=1):
        safe_name = os.path.basename(file.filename or f"resume_{index}.pdf")
        ext = get_file_extension(safe_name)

        if ext not in ALLOWED_RESUME_EXTENSIONS:
            skipped_files.append(safe_name)
            continue

        file_path = os.path.join(HYBRID_UPLOAD_DIR, safe_name)
        content = await file.read()

        if not content:
            skipped_files.append(safe_name)
            continue

        with open(file_path, "wb") as f:
            f.write(content)

        saved_files.append(
            {
                "filename": safe_name,
                "content_type": file.content_type or "application/pdf",
                "path": file_path,
            }
        )

    if not saved_files:
        raise HTTPException(
            status_code=400,
            detail="No supported resume files found. Supported: PDF, DOCX, DOC.",
        )

    status = get_default_status()
    status.update(
        {
            "job_id": job_id,
            "status": "queued",
            "message": "Hybrid resume analysis job queued",
            "total": len(saved_files),
            "processed": 0,
            "successful": 0,
            "failed": 0,
            "skipped": len(skipped_files),
            "skipped_files": skipped_files,
            "current_file": "",
            "current_batch": f"0/{len(saved_files)}",
            "started_at": datetime.now().isoformat(),
        }
    )

    write_status(status)

    asyncio.create_task(run_hybrid_background_job(job_id, saved_files, cfg))

    return {
        "job_id": job_id,
        "status": "queued",
        "total": len(saved_files),
        "message": "Hybrid resume analysis started",
    }


@app.get("/hybrid-status")
def get_hybrid_status():
    return read_status()


@app.post("/hybrid-reset-status")
def reset_hybrid_status():
    if is_active_job_running():
        raise HTTPException(
            status_code=409,
            detail="Cannot reset while hybrid resume analysis is running.",
        )

    reset_hybrid_job_storage()

    status = get_default_status()
    write_status(status)

    return {
        "success": True,
        "message": "Hybrid status reset successfully.",
    }


@app.get("/hybrid-download/{filename}")
def hybrid_download(filename: str):
    safe_filename = os.path.basename(filename)
    path = os.path.join(OUTPUT_DIR, safe_filename)

    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Hybrid output file not found.")

    return FileResponse(
        path,
        filename=safe_filename,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
