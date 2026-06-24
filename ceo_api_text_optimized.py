import os
import re
import csv
import time
import uuid
import json
import shutil
import threading
from pathlib import Path
from typing import Any, Dict, List

import fitz  # pip install PyMuPDF
from docx import Document  # pip install python-docx

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS

# =========================================================
# IMPORT CEO ORIGINAL LOGIC
# =========================================================
# This keeps Excel output EXACTLY same as CEO original app.
from ceo_original_app import (
    run_match_job,
    job_state,
    _state_snapshot,
    OUTPUT_DIR,
    WORK_DIR,
    token_tracker,
)

# =========================================================
# APP CONFIG
# =========================================================

app = Flask(__name__)
CORS(app)

PORT = 8008
PIPELINE_NAME = "ceo_text_optimized_same_output"

CEO_API_WORK_DIR = WORK_DIR / "api_text_optimized"
RAW_RESUME_DIR = CEO_API_WORK_DIR / "raw_resumes"
TEXT_RESUME_DIR = CEO_API_WORK_DIR / "text_resumes"
LOG_DIR = CEO_API_WORK_DIR / "logs"

REQUIREMENT_EXCEL_PATH = CEO_API_WORK_DIR / "openings.xlsx"

CONVERSION_LOG_JSON = LOG_DIR / "conversion_log.json"
COST_LOG_JSON = LOG_DIR / "cost_log.json"
COST_LOG_CSV = LOG_DIR / "cost_log.csv"

for folder in [CEO_API_WORK_DIR, RAW_RESUME_DIR, TEXT_RESUME_DIR, LOG_DIR]:
    folder.mkdir(parents=True, exist_ok=True)

# =========================================================
# COST CONFIG
# =========================================================

USD_TO_INR = float(os.getenv("USD_TO_INR", "83.0"))

# Update this based on your Claude model pricing.
INPUT_PRICE_PER_1M_TOKENS_USD = float(os.getenv("INPUT_PRICE_PER_1M_TOKENS_USD", "3.0"))
OUTPUT_PRICE_PER_1M_TOKENS_USD = float(os.getenv("OUTPUT_PRICE_PER_1M_TOKENS_USD", "15.0"))

MAX_COST_PER_RESUME_INR = float(os.getenv("MAX_COST_PER_RESUME_INR", "0.30"))
CHARS_PER_TOKEN_ESTIMATE = float(os.getenv("CHARS_PER_TOKEN_ESTIMATE", "4.0"))

# Do not block by default. Only show warning in UI.
HARD_BLOCK_OVER_BUDGET_RESUME = False

text_mode_metrics: Dict[str, Any] = {
    "job_id": None,
    "pipeline": PIPELINE_NAME,
    "started_at": None,
    "finished_at": None,
    "total_uploaded": 0,
    "total_raw_saved": 0,
    "total_converted_to_text": 0,
    "total_skipped": 0,
    "conversion_rows": [],
    "cost_rows": [],
    "summary": {},
}


# =========================================================
# HELPERS
# =========================================================

def safe_filename(filename: str) -> str:
    filename = os.path.basename(filename or "")
    filename = re.sub(r"[^a-zA-Z0-9_.\-\[\]\(\) ]", "_", filename)
    filename = re.sub(r"\s+", " ", filename).strip()
    return filename or f"file_{uuid.uuid4().hex[:8]}"


def append_log(message: str):
    try:
        job_state.setdefault("log", [])
        job_state["log"].append({
            "time": time.strftime("%H:%M:%S"),
            "message": message,
        })
    except Exception:
        pass


def save_json(path: Path, data: Any):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def read_json(path: Path, default: Any):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def write_cost_csv(rows: List[Dict[str, Any]]):
    if not rows:
        COST_LOG_CSV.write_text("", encoding="utf-8")
        return

    fieldnames = list(rows[0].keys())

    with COST_LOG_CSV.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def clean_line(line: str) -> str:
    if not line:
        return ""

    line = line.replace("\r", "")
    line = line.replace("\t", " ")
    line = re.sub(r"[ ]{2,}", " ", line)

    return line.strip()


def estimate_tokens(text: str) -> int:
    if not text:
        return 0
    return max(1, int(len(text) / CHARS_PER_TOKEN_ESTIMATE))


def estimate_input_cost(input_tokens: int) -> Dict[str, float]:
    usd = (input_tokens / 1_000_000) * INPUT_PRICE_PER_1M_TOKENS_USD
    inr = usd * USD_TO_INR

    return {
        "estimated_input_cost_usd": round(usd, 6),
        "estimated_input_cost_inr": round(inr, 4),
    }


def reset_storage():
    for folder in [RAW_RESUME_DIR, TEXT_RESUME_DIR, LOG_DIR]:
        if folder.exists():
            shutil.rmtree(folder)
        folder.mkdir(parents=True, exist_ok=True)

    if REQUIREMENT_EXCEL_PATH.exists():
        REQUIREMENT_EXCEL_PATH.unlink()

    save_json(CONVERSION_LOG_JSON, [])
    save_json(COST_LOG_JSON, [])
    COST_LOG_CSV.write_text("", encoding="utf-8")


def reset_ceo_job_state():
    job_state.update({
        "status": "idle",
        "stage": "",
        "progress": 0,
        "total": 0,
        "message": "",
        "log": [],
        "outputs": {},
        "results_summary": None,
        "start_time": None,
    })

    try:
        token_tracker.reset()
    except Exception:
        pass


def reset_text_mode_metrics():
    text_mode_metrics.update({
        "job_id": None,
        "pipeline": PIPELINE_NAME,
        "started_at": None,
        "finished_at": None,
        "total_uploaded": 0,
        "total_raw_saved": 0,
        "total_converted_to_text": 0,
        "total_skipped": 0,
        "conversion_rows": [],
        "cost_rows": [],
        "summary": {},
    })


# =========================================================
# TEXT EXTRACTION
# =========================================================

def extract_pdf_to_line_text(pdf_path: Path) -> str:
    doc = fitz.open(str(pdf_path))

    output_lines = [
        f"SOURCE_FILE: {pdf_path.name}",
        "SOURCE_TYPE: PDF",
        "EXTRACTION_MODE: exact_line_by_line_text",
        "",
        "IMPORTANT_NOTE:",
        "This is exact extracted text from resume. Line numbers are added only for evidence.",
        "Do not treat Date of Birth as job duration.",
        "Do not treat education/client/company/project location as candidate current location.",
        "",
    ]

    total_lines = 0

    for page_no, page in enumerate(doc, start=1):
        output_lines.append(f"========== PAGE {page_no} START ==========")

        raw_text = page.get_text("text", sort=True)
        page_line_no = 0

        for raw_line in raw_text.splitlines():
            line = clean_line(raw_line)

            if not line:
                continue

            page_line_no += 1
            total_lines += 1

            output_lines.append(f"[PAGE {page_no} LINE {page_line_no:03d}] {line}")

        output_lines.append(f"========== PAGE {page_no} END ==========")
        output_lines.append("")

    doc.close()

    output_lines.append(f"TOTAL_EXTRACTED_LINES: {total_lines}")

    return "\n".join(output_lines).strip()


def extract_docx_to_line_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))

    output_lines = [
        f"SOURCE_FILE: {docx_path.name}",
        "SOURCE_TYPE: DOCX",
        "EXTRACTION_MODE: exact_line_by_line_text",
        "",
        "IMPORTANT_NOTE:",
        "This is exact extracted text from resume. Line numbers are added only for evidence.",
        "Do not treat Date of Birth as job duration.",
        "Do not treat education/client/company/project location as candidate current location.",
        "",
        "========== DOCUMENT PARAGRAPHS START ==========",
    ]

    line_no = 0

    for para in doc.paragraphs:
        line = clean_line(para.text)

        if not line:
            continue

        line_no += 1
        output_lines.append(f"[DOC LINE {line_no:03d}] {line}")

    output_lines.append("========== DOCUMENT PARAGRAPHS END ==========")
    output_lines.append("")

    if doc.tables:
        output_lines.append("========== DOCUMENT TABLES START ==========")

        for table_no, table in enumerate(doc.tables, start=1):
            output_lines.append(f"----- TABLE {table_no} START -----")

            for row_no, row in enumerate(table.rows, start=1):
                cells = []

                for cell in row.cells:
                    cell_text = clean_line(cell.text.replace("\n", " "))

                    if cell_text:
                        cells.append(cell_text)

                if not cells:
                    continue

                line_no += 1
                output_lines.append(
                    f"[DOC LINE {line_no:03d}] "
                    f"[TABLE {table_no} ROW {row_no}] "
                    + " | ".join(cells)
                )

            output_lines.append(f"----- TABLE {table_no} END -----")
            output_lines.append("")

        output_lines.append("========== DOCUMENT TABLES END ==========")

    output_lines.append(f"TOTAL_EXTRACTED_LINES: {line_no}")

    return "\n".join(output_lines).strip()


def convert_resume_to_text_file(raw_path: Path) -> Dict[str, Any]:
    ext = raw_path.suffix.lower()

    if ext == ".pdf":
        extracted_text = extract_pdf_to_line_text(raw_path)
    elif ext == ".docx":
        extracted_text = extract_docx_to_line_text(raw_path)
    elif ext == ".txt":
        extracted_text = raw_path.read_text(encoding="utf-8", errors="ignore")
    elif ext == ".doc":
        raise ValueError(
            "Old .doc is not supported in text optimized mode. "
            "Please convert .doc to .docx before upload."
        )
    else:
        raise ValueError(f"Unsupported resume file type: {ext}")

    txt_name = f"{raw_path.stem}.txt"
    txt_path = TEXT_RESUME_DIR / txt_name
    txt_path.write_text(extracted_text, encoding="utf-8")

    input_tokens = estimate_tokens(extracted_text)
    input_cost = estimate_input_cost(input_tokens)
    over_budget = input_cost["estimated_input_cost_inr"] > MAX_COST_PER_RESUME_INR

    conversion_row = {
        "original_filename": raw_path.name,
        "text_filename": txt_name,
        "raw_path": str(raw_path),
        "text_path": str(txt_path),
        "status": "converted",
        "error": "",
        "characters": len(extracted_text),
        "estimated_input_tokens": input_tokens,
        "estimated_input_cost_inr": input_cost["estimated_input_cost_inr"],
        "over_budget_estimate": over_budget,
    }

    cost_row = {
        "original_filename": raw_path.name,
        "text_filename": txt_name,
        "status": "converted",
        "characters": len(extracted_text),
        "estimated_input_tokens": input_tokens,
        "estimated_input_cost_usd": input_cost["estimated_input_cost_usd"],
        "estimated_input_cost_inr": input_cost["estimated_input_cost_inr"],
        "max_cost_per_resume_inr": MAX_COST_PER_RESUME_INR,
        "over_budget_estimate": over_budget,
        "note": "Estimated resume text input cost only. Actual total cost depends on prompt, requirements, output tokens, and CEO original calls.",
    }

    return {
        "ok": True,
        "txt_path": txt_path,
        "conversion_row": conversion_row,
        "cost_row": cost_row,
    }


# =========================================================
# MONITOR
# =========================================================

def monitor_job(job_id: str):
    while True:
        status = job_state.get("status")

        if status in {"completed", "done", "failed", "error", "idle"}:
            break

        time.sleep(2)

    text_mode_metrics["finished_at"] = time.time()

    try:
        actual_token_snapshot = token_tracker.snapshot()
    except Exception as e:
        actual_token_snapshot = {"error": str(e)}

    total_est_tokens = sum(
        row.get("estimated_input_tokens", 0)
        for row in text_mode_metrics.get("cost_rows", [])
    )

    total_est_cost_inr = sum(
        row.get("estimated_input_cost_inr", 0.0)
        for row in text_mode_metrics.get("cost_rows", [])
    )

    over_budget_count = sum(
        1 for row in text_mode_metrics.get("cost_rows", [])
        if row.get("over_budget_estimate")
    )

    elapsed_seconds = None
    if text_mode_metrics.get("started_at"):
        elapsed_seconds = round(time.time() - text_mode_metrics["started_at"], 2)

    text_mode_metrics["summary"] = {
        "job_id": job_id,
        "pipeline": PIPELINE_NAME,
        "job_status": job_state.get("status"),
        "total_uploaded": text_mode_metrics.get("total_uploaded", 0),
        "total_raw_saved": text_mode_metrics.get("total_raw_saved", 0),
        "total_converted_to_text": text_mode_metrics.get("total_converted_to_text", 0),
        "total_skipped": text_mode_metrics.get("total_skipped", 0),
        "total_estimated_input_tokens_from_text_files": total_est_tokens,
        "total_estimated_input_cost_inr_from_text_files": round(total_est_cost_inr, 4),
        "max_cost_per_resume_inr": MAX_COST_PER_RESUME_INR,
        "over_budget_resume_count": over_budget_count,
        "elapsed_seconds": elapsed_seconds,
        "actual_token_tracker_snapshot_from_ceo_original": actual_token_snapshot,
        "output_format_note": "Final Excel output is generated by CEO original run_match_job, so output format remains same.",
    }

    save_json(CONVERSION_LOG_JSON, text_mode_metrics["conversion_rows"])
    save_json(COST_LOG_JSON, text_mode_metrics["cost_rows"])
    write_cost_csv(text_mode_metrics["cost_rows"])

    append_log("Text optimized cost monitor completed.")
    append_log(f"Estimated text input cost INR: {round(total_est_cost_inr, 4)}")


# =========================================================
# ROUTES
# =========================================================

@app.get("/ceo-health")
def ceo_health():
    return jsonify({
        "ok": True,
        "pipeline": PIPELINE_NAME,
        "message": "CEO Text Optimized API running",
        "port": PORT,
        "output_format": "same_as_ceo_original_main_app",
        "raw_resume_dir": str(RAW_RESUME_DIR),
        "text_resume_dir": str(TEXT_RESUME_DIR),
        "log_dir": str(LOG_DIR),
    })


@app.post("/ceo-start")
def ceo_start():
    if job_state.get("status") == "running":
        return jsonify({"ok": False, "error": "CEO job already running"}), 409

    api_key = request.form.get("api_key", "").strip()
    excel_file = request.files.get("excel")
    resume_files = request.files.getlist("files")

    if not api_key.startswith("sk-ant"):
        return jsonify({"ok": False, "error": "Invalid or missing Anthropic API key"}), 400

    if not excel_file:
        return jsonify({"ok": False, "error": "Requirement Excel file is missing"}), 400

    if not resume_files:
        return jsonify({"ok": False, "error": "Resume files are missing"}), 400

    reset_storage()
    reset_ceo_job_state()
    reset_text_mode_metrics()

    job_id = f"CEO_TEXT_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"

    text_mode_metrics["job_id"] = job_id
    text_mode_metrics["started_at"] = time.time()
    text_mode_metrics["total_uploaded"] = len(resume_files)

    excel_file.save(REQUIREMENT_EXCEL_PATH)

    allowed_ext = {".pdf", ".docx", ".txt", ".doc"}

    raw_saved_count = 0
    converted_count = 0
    skipped_files = []
    conversion_errors = []
    conversion_rows = []
    cost_rows = []

    for uploaded_file in resume_files:
        original_name = safe_filename(uploaded_file.filename)
        ext = Path(original_name).suffix.lower()

        if ext not in allowed_ext:
            skipped_files.append({"filename": original_name, "reason": "Unsupported extension"})
            continue

        raw_path = RAW_RESUME_DIR / original_name
        uploaded_file.save(raw_path)
        raw_saved_count += 1

        try:
            result = convert_resume_to_text_file(raw_path)
            conversion_row = result["conversion_row"]
            cost_row = result["cost_row"]

            if HARD_BLOCK_OVER_BUDGET_RESUME and cost_row["over_budget_estimate"]:
                skipped_files.append({
                    "filename": original_name,
                    "reason": f"Over estimated budget �{MAX_COST_PER_RESUME_INR}",
                })
                try:
                    result["txt_path"].unlink()
                except Exception:
                    pass
                continue

            conversion_rows.append(conversion_row)
            cost_rows.append(cost_row)
            converted_count += 1

            append_log(f"Converted to TXT: {original_name} -> {conversion_row['text_filename']}")

        except Exception as e:
            err = str(e)
            conversion_errors.append({"filename": original_name, "error": err})

            conversion_rows.append({
                "original_filename": original_name,
                "text_filename": "",
                "raw_path": str(raw_path),
                "text_path": "",
                "status": "failed",
                "error": err,
                "characters": 0,
                "estimated_input_tokens": 0,
                "estimated_input_cost_inr": 0,
                "over_budget_estimate": False,
            })

            cost_rows.append({
                "original_filename": original_name,
                "text_filename": "",
                "status": "failed",
                "characters": 0,
                "estimated_input_tokens": 0,
                "estimated_input_cost_usd": 0,
                "estimated_input_cost_inr": 0,
                "max_cost_per_resume_inr": MAX_COST_PER_RESUME_INR,
                "over_budget_estimate": False,
                "note": err,
            })

    text_mode_metrics["total_raw_saved"] = raw_saved_count
    text_mode_metrics["total_converted_to_text"] = converted_count
    text_mode_metrics["total_skipped"] = len(skipped_files) + len(conversion_errors)
    text_mode_metrics["conversion_rows"] = conversion_rows
    text_mode_metrics["cost_rows"] = cost_rows

    save_json(CONVERSION_LOG_JSON, conversion_rows)
    save_json(COST_LOG_JSON, cost_rows)
    write_cost_csv(cost_rows)

    if converted_count == 0:
        return jsonify({
            "ok": False,
            "error": "No resume could be converted to text.",
            "skipped_files": skipped_files,
            "conversion_errors": conversion_errors,
        }), 400

    total_est_tokens = sum(row["estimated_input_tokens"] for row in cost_rows)
    total_est_cost_inr = round(sum(row["estimated_input_cost_inr"] for row in cost_rows), 4)

    job_state.update({
        "status": "running",
        "stage": "text_conversion_done",
        "progress": 0,
        "total": converted_count,
        "message": "CEO matching started with text optimized resumes",
        "log": job_state.get("log", []),
        "outputs": {},
        "results_summary": None,
        "start_time": time.time(),
        "job_id": job_id,
        "pipeline": PIPELINE_NAME,
    })

    try:
        token_tracker.reset()
    except Exception:
        pass

    append_log("All supported resumes converted to exact line-by-line TXT.")
    append_log(f"Total TXT resumes: {converted_count}")
    append_log("Starting CEO original run_match_job() using TXT resumes.")
    append_log("Final Excel output format remains same as CEO original app.")

    # IMPORTANT:
    # Passing TEXT_RESUME_DIR keeps cost lower.
    # run_match_job still creates original CEO Excel output.
    threading.Thread(
        target=run_match_job,
        args=(api_key, str(REQUIREMENT_EXCEL_PATH), str(TEXT_RESUME_DIR)),
        daemon=True,
    ).start()

    threading.Thread(
        target=monitor_job,
        args=(job_id,),
        daemon=True,
    ).start()

    return jsonify({
        "ok": True,
        "job_id": job_id,
        "pipeline": PIPELINE_NAME,
        "message": "CEO Text Optimized pipeline started.",
        "output_format": "same_as_ceo_original_main_app",
        "total_files_received": len(resume_files),
        "total_raw_saved": raw_saved_count,
        "total_text_converted": converted_count,
        "skipped_files": skipped_files,
        "conversion_errors": conversion_errors,
        "preflight_cost_summary": {
            "estimated_input_tokens_from_text_files": total_est_tokens,
            "estimated_input_cost_inr_from_text_files": total_est_cost_inr,
            "max_cost_per_resume_inr": MAX_COST_PER_RESUME_INR,
            "note": "Estimated resume text input cost only. Actual cost depends on CEO original prompt, requirements, output tokens, and model calls.",
        },
    })


@app.get("/ceo-status")
def ceo_status():
    snapshot = _state_snapshot()

    outputs = snapshot.get("outputs", {}) or {}

    if outputs.get("excel"):
        outputs["excel_download_url"] = f"/ceo-download/{outputs['excel']}"

    if outputs.get("pdf"):
        outputs["pdf_download_url"] = f"/ceo-download/{outputs['pdf']}"

    recent_log = snapshot.get("recent_log")
    if not recent_log:
        log_items = snapshot.get("log") or job_state.get("log") or []
        recent_log = log_items[-20:]

    snapshot["outputs"] = outputs
    snapshot["pipeline"] = PIPELINE_NAME
    snapshot["output_format"] = "same_as_ceo_original_main_app"
    snapshot["text_mode_summary"] = text_mode_metrics.get("summary", {})
    snapshot["text_converted_count"] = text_mode_metrics.get("total_converted_to_text", 0)
    snapshot["text_skipped_count"] = text_mode_metrics.get("total_skipped", 0)
    snapshot["recent_log"] = recent_log

    return jsonify(snapshot)


@app.get("/ceo-token-cost")
def ceo_token_cost():
    try:
        actual_snapshot = token_tracker.snapshot()
    except Exception as e:
        actual_snapshot = {"error": str(e)}

    return jsonify({
        "ok": True,
        "pipeline": PIPELINE_NAME,
        "actual_token_tracker_from_ceo_original": actual_snapshot,
        "text_mode_summary": text_mode_metrics.get("summary", {}),
    })


@app.get("/ceo-cost-log")
def ceo_cost_log():
    rows = read_json(COST_LOG_JSON, [])

    total_tokens = sum(row.get("estimated_input_tokens", 0) for row in rows)
    total_cost_inr = round(sum(row.get("estimated_input_cost_inr", 0.0) for row in rows), 4)

    return jsonify({
        "ok": True,
        "pipeline": PIPELINE_NAME,
        "rows": rows,
        "summary": {
            "total_resumes": len(rows),
            "total_estimated_input_tokens_from_text_files": total_tokens,
            "total_estimated_input_cost_inr_from_text_files": total_cost_inr,
            "max_cost_per_resume_inr": MAX_COST_PER_RESUME_INR,
            "csv_download_url": "/ceo-cost-log-csv",
        },
    })


@app.get("/ceo-cost-log-csv")
def ceo_cost_log_csv():
    if not COST_LOG_CSV.exists():
        return jsonify({"ok": False, "error": "Cost CSV not found"}), 404

    return send_from_directory(LOG_DIR, COST_LOG_CSV.name, as_attachment=True)


@app.get("/ceo-conversion-log")
def ceo_conversion_log():
    return jsonify({
        "ok": True,
        "pipeline": PIPELINE_NAME,
        "rows": read_json(CONVERSION_LOG_JSON, []),
    })


@app.get("/ceo-converted-text-files")
def ceo_converted_text_files():
    files = []

    for p in TEXT_RESUME_DIR.glob("*.txt"):
        files.append({
            "filename": p.name,
            "size_bytes": p.stat().st_size,
            "download_url": f"/ceo-text-download/{p.name}",
        })

    return jsonify({
        "ok": True,
        "pipeline": PIPELINE_NAME,
        "files": files,
    })


@app.get("/ceo-text-download/<filename>")
def ceo_text_download(filename):
    safe = safe_filename(filename)
    path = TEXT_RESUME_DIR / safe

    if not path.exists():
        return jsonify({"ok": False, "error": "Converted text file not found"}), 404

    return send_from_directory(TEXT_RESUME_DIR, safe, as_attachment=True)


@app.get("/ceo-dashboard-data")
def ceo_dashboard_data():
    return jsonify(job_state.get("outputs", {}).get("dashboard_data", []))


@app.get("/ceo-download/<filename>")
def ceo_download(filename):
    safe = re.sub(r"[^a-zA-Z0-9_.-]", "", filename)
    file_path = Path(OUTPUT_DIR) / safe

    if not file_path.exists():
        return jsonify({"ok": False, "error": "File not found"}), 404

    return send_from_directory(OUTPUT_DIR, safe, as_attachment=True)


@app.post("/ceo-reset")
def ceo_reset():
    if job_state.get("status") == "running":
        return jsonify({"ok": False, "error": "Cannot reset while job is running"}), 409

    reset_storage()
    reset_ceo_job_state()
    reset_text_mode_metrics()

    return jsonify({
        "ok": True,
        "message": "CEO Text Optimized API state reset successfully",
    })


if __name__ == "__main__":
    print(f"\nCEO Text Optimized API running on http://localhost:{PORT}")
    print("Output Excel format: SAME as CEO original app")
    print("Resume input to Claude: TXT only, not raw PDF/DOCX")
    print("\nRoutes:")
    print("GET  /ceo-health")
    print("POST /ceo-start")
    print("GET  /ceo-status")
    print("GET  /ceo-token-cost")
    print("GET  /ceo-cost-log")
    print("GET  /ceo-cost-log-csv")
    print("GET  /ceo-conversion-log")
    print("GET  /ceo-converted-text-files")
    print("GET  /ceo-text-download/<filename>")
    print("GET  /ceo-download/<filename>")
    print("POST /ceo-reset\n")

    app.run(host="0.0.0.0", port=PORT, debug=False)